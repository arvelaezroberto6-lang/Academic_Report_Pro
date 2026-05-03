from flask import Flask, render_template, request, jsonify, send_file
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch, cm
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, HRFlowable, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import cm as rl_cm
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import os
import uuid
from datetime import datetime
import requests as http_requests
import logging
import re
import html
import time
from concurrent.futures import ThreadPoolExecutor, as_completed

# ── Importar módulo de referencias reales ──────────────────────
from referencias_reales import buscar_referencias_reales, formatear_referencias

# ── Importar funciones de base de datos ────────────────────────
from database import (
    registrar_usuario, login_usuario, obtener_perfil,
    actualizar_perfil, guardar_informe, obtener_mis_informes,
    obtener_informe, eliminar_informe, DB_DISPONIBLE,
    obtener_estadisticas_usuario,
    obtener_resumen_actividad,
)


# ── Importar módulo de seguridad ───────────────────────────────
from security import (
    requiere_auth, obtener_user_id_verificado,
    sanitizar_texto, sanitizar_html_email, sanitizar_nombre,
    es_uuid_valido, es_email_valido, es_password_seguro,
    validar_params_informe, aplicar_headers_seguridad,
    log_evento_seguridad, NORMAS_VALIDAS, TIPOS_VALIDOS,
    NIVELES_VALIDOS, MODOS_VALIDOS
)

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 2 * 1024 * 1024  # máx 2 MB por request
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'cambia-esto-en-produccion')

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ── Headers de seguridad en todas las respuestas ──────────────
app.after_request(aplicar_headers_seguridad)

# ── Rate Limiter (requiere: pip install flask-limiter) ─────────
try:
    from flask_limiter import Limiter
    from flask_limiter.util import get_remote_address
    limiter = Limiter(
        get_remote_address,
        app=app,
        default_limits=["500 per day", "100 per hour"],
        storage_uri="memory://",   # cambiar a Redis en producción con: redis://localhost:6379
        headers_enabled=True
    )
    LIMITER_DISPONIBLE = True
    logger.info("✅ Rate limiter activo")
except ImportError:
    limiter = None
    LIMITER_DISPONIBLE = False
    logger.warning("⚠️ flask-limiter no instalado — sin protección de rate limit. Ejecuta: pip install flask-limiter")

def limit(rule):
    """Decorador de rate limit que falla silencioso si no hay limiter."""
    def decorator(f):
        if LIMITER_DISPONIBLE and limiter:
            return limiter.limit(rule)(f)
        return f
    return decorator

os.makedirs('informes_generados', exist_ok=True)

DEEPSEEK_API_KEY = os.environ.get('DEEPSEEK_API_KEY', '')
DEEPSEEK_URL     = "https://api.deepseek.com/v1/chat/completions"

logger.info("=" * 60)
logger.info("🚀 ACADEMIC REPORT PRO - VERSIÓN 3.3")
logger.info(f"🔑 API Key: {'SÍ ✅' if DEEPSEEK_API_KEY else 'NO ❌'}")
logger.info("=" * 60)

# ============================================================
# NORMAS Y TIPOS
# ============================================================
NORMAS_INSTRUCCIONES = {
    'APA 7': """NORMA APA 7 — En texto: (Apellido, año). 3+ autores: primer apellido et al.
Referencias: Apellido, I. (año). Título. Revista, vol(num), págs. https://doi.org/...""",
    'APA 6': """NORMA APA 6 — En texto: (Apellido, año, p. XX). 3-5 autores: primera vez todos; luego et al.
Referencias: Apellido, I. (año). Título del libro en cursiva. Ciudad: Editorial.""",
    'ICONTEC': """NORMA ICONTEC (NTC 5613) — Notas al pie numeradas. Referencias en ORDEN DE APARICIÓN.
Formato: APELLIDO, Nombre. Título en cursiva. Ed. Ciudad: Editorial, año.""",
    'IEEE': """NORMA IEEE — En texto: [1], [2]. Referencias numeradas en orden de aparición.
Artículos: [1] I. Apellido, "Título," Revista, vol. X, no. X, pp. XX, año.""",
    'Vancouver': """NORMA VANCOUVER — Números superíndice en orden de aparición.
Artículos: Apellido AB. Título. Revista. año;vol(num):págs.""",
    'Chicago': """NORMA CHICAGO 17 — En texto: (Apellido año, página). Bibliografía alfabética.
Libros: Apellido, Nombre. Año. Título. Ciudad: Editorial.""",
    'MLA': """NORMA MLA 9 — En texto: (Apellido página). Works Cited alfabético.
Artículos: Apellido, Nombre. "Título." Revista, vol. X, no. X, año, pp. XX.""",
    'Harvard': """NORMA HARVARD — En texto: (Apellido año). Referencias alfabéticas.
Artículos: Apellido, I. (año) 'Título', Revista, vol. X, no. X, pp. XX.""",
}

TIPOS_INSTRUCCIONES = {
    'academico':  "Informe académico estándar con rigor teórico, citas y análisis crítico.",
    'laboratorio':"Informe de laboratorio: hipótesis, materiales, procedimiento, resultados con datos/tablas, análisis de error.",
    'ejecutivo':  "Informe ejecutivo: lenguaje conciso, KPIs, análisis costo-beneficio, recomendaciones accionables.",
    'tesis':      "Tesis académica: argumentación profunda, revisión exhaustiva de literatura, marco metodológico robusto.",
    'pasantia':   "Informe de pasantía/práctica empresarial: describe la organización, el rol del practicante, actividades realizadas por etapas, competencias desarrolladas, herramientas usadas y lecciones aprendidas. Tono reflexivo y profesional.",
    'proyecto':   "Informe de proyecto (gestión de proyectos): describe alcance, fases, recursos, entregables, gestión de riesgos, indicadores de éxito (KPIs), estado de avance y lecciones aprendidas. Estructura orientada a resultados medibles.",
}

# ============================================================
# DEEPSEEK
# ============================================================
def llamar_deepseek(prompt, system_prompt=None, max_tokens=3000, reintentos=3):
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }
    messages = []
    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})
    messages.append({"role": "user", "content": prompt})
    data = {
        "model":       "deepseek-chat",
        "messages":    messages,
        "max_tokens":  max_tokens,
        "temperature": 0.7
    }
    for intento in range(1, reintentos + 1):
        try:
            response = http_requests.post(DEEPSEEK_URL, headers=headers, json=data, timeout=150)
            if response.status_code == 200:
                contenido = response.json()['choices'][0]['message']['content']
                return contenido.encode('utf-8', 'ignore').decode('utf-8')
            elif response.status_code in (429, 503, 502):
                wait = 5 * intento
                logger.warning(f"DeepSeek HTTP {response.status_code} — reintento {intento}/{reintentos} en {wait}s")
                time.sleep(wait)
            else:
                logger.error(f"DeepSeek HTTP {response.status_code}: {response.text[:200]}")
                return None
        except http_requests.exceptions.Timeout:
            wait = 8 * intento
            logger.warning(f"Timeout DeepSeek — reintento {intento}/{reintentos} en {wait}s")
            time.sleep(wait)
        except Exception as e:
            logger.error(f"Error DeepSeek: {e}")
            if intento < reintentos:
                time.sleep(5 * intento)
            else:
                return None
    logger.error(f"DeepSeek falló después de {reintentos} reintentos")
    return None

# ============================================================
# LIMPIEZA DE TEXTO
# ============================================================
def limpiar_para_pdf(texto):
    if not texto:
        return ""
    texto = texto.replace('<br/>', '\n').replace('<br>', '\n')
    texto = re.sub(r'<[^>]+>', '', texto)
    texto = html.unescape(texto)
    texto = re.sub(r'\n{3,}', '\n\n', texto)
    return texto.strip()

def limpiar_para_word(texto):
    if not texto:
        return ""
    texto = texto.replace('<br/>', '\n').replace('<br>', '\n')
    texto = re.sub(r'<[^>]+>', '', texto)
    texto = html.unescape(texto)
    return texto.strip()

# ============================================================
# VALIDADORES DE CALIDAD (post-generación)
# ============================================================
_PATRONES_CITA = {
    'APA 7':    re.compile(r'\([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+.*?,\s*20[1-2]\d\)'),
    'APA 6':    re.compile(r'\([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+.*?,\s*20[1-2]\d\)'),
    'ICONTEC':  re.compile(r'(?:\d+\s*\)|\[?\d+\]?\s*[A-ZÁÉÍÓÚÑ])'),
    'IEEE':     re.compile(r'\[\d+\]'),
    'Vancouver':re.compile(r'(?:^\d+\.|(?<!\w)\d{1,2}(?!\w))'),
    'Chicago':  re.compile(r'\([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+\s+20[1-2]\d\)'),
    'MLA':      re.compile(r'\([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+\s+\d+\)'),
    'Harvard':  re.compile(r'\([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+.*?20[1-2]\d\)'),
}
_SECCIONES_CON_CITAS    = {'introduccion', 'marco_teorico', 'desarrollo', 'metodologia'}
_MIN_CITAS_POR_SECCION  = {'introduccion': 2, 'marco_teorico': 4, 'desarrollo': 4, 'metodologia': 1}

def validar_citas(contenido: str, seccion: str, norma: str) -> dict:
    if seccion not in _SECCIONES_CON_CITAS or not contenido:
        return {'ok': True, 'citas_encontradas': 0, 'minimo': 0}
    patron = _PATRONES_CITA.get(norma, _PATRONES_CITA['APA 7'])
    citas  = patron.findall(contenido)
    minimo = _MIN_CITAS_POR_SECCION.get(seccion, 1)
    return {'ok': len(citas) >= minimo, 'citas_encontradas': len(citas), 'minimo': minimo}

def validar_tabla_en_desarrollo(contenido: str) -> bool:
    tiene_struct = '##TABLE##' in contenido and '##ENDTABLE##' in contenido
    tiene_md     = bool(re.search(r'^\|.+\|', contenido, re.MULTILINE))
    return tiene_struct or tiene_md

# ============================================================
# PARSEO DE TABLAS ESTRUCTURADAS
# ============================================================
def extraer_tablas(texto: str):
    """
    Detecta bloques ##TABLE## ... ##ENDTABLE## y también tablas markdown (| col | col |).
    Devuelve lista de dicts: {titulo, cabeceras: [], filas: [[]], inicio, fin}
    y el texto con los bloques reemplazados por marcadores únicos.
    """
    tablas = []

    # ── Formato estructurado ##TABLE## ──
    patron = re.compile(r'##TABLE##(.*?)##ENDTABLE##', re.DOTALL)
    def reemplazar(m):
        bloque = m.group(1).strip()
        titulo = ""
        cabeceras = []
        filas = []
        for linea in bloque.splitlines():
            linea = linea.strip()
            if linea.startswith("TITULO:"):
                titulo = linea[7:].strip()
            elif linea.startswith("CABECERAS:"):
                cabeceras = [c.strip() for c in linea[10:].split("|")]
            elif linea.startswith("FILA:"):
                fila = [c.strip() for c in linea[5:].split("|")]
                if fila:
                    filas.append(fila)
        if cabeceras and filas:
            idx = len(tablas)
            tablas.append({"titulo": titulo, "cabeceras": cabeceras, "filas": filas})
            return f"__TABLA_{idx}__"
        return bloque  # si malformado, dejar como texto
    texto = patron.sub(reemplazar, texto)

    # ── Formato markdown  | col | col | ──
    lineas = texto.splitlines()
    i = 0
    nuevo_texto = []
    while i < len(lineas):
        linea = lineas[i].strip()
        if linea.startswith("|") and linea.endswith("|") and linea.count("|") >= 3:
            # Detectar bloque de tabla
            bloque_md = []
            while i < len(lineas) and lineas[i].strip().startswith("|"):
                bloque_md.append(lineas[i].strip())
                i += 1
            # Filtrar separadores (|---|---|)
            filas_md = [l for l in bloque_md if not re.match(r'^\|[-:\s|]+\|$', l)]
            parsed = [[c.strip() for c in f.strip("|").split("|")] for f in filas_md]
            if len(parsed) >= 2:
                # Buscar si hay un titulo antes
                titulo_md = ""
                if nuevo_texto:
                    ultima = nuevo_texto[-1].strip()
                    if ultima.lower().startswith("tabla") and len(ultima) < 120:
                        titulo_md = nuevo_texto.pop()
                idx = len(tablas)
                tablas.append({
                    "titulo": titulo_md,
                    "cabeceras": parsed[0],
                    "filas": parsed[1:]
                })
                nuevo_texto.append(f"__TABLA_{idx}__")
            else:
                nuevo_texto.extend(bloque_md)
            continue
        nuevo_texto.append(lineas[i])
        i += 1
    texto = "\n".join(nuevo_texto)
    return texto, tablas


# ============================================================
# PROMPTS POR SECCIÓN  (v3.3 — mejorado con recomendaciones)
# ============================================================
# Contexto colombiano inyectado en todas las secciones relevantes
_CONTEXTO_COLOMBIA = """
CONTEXTO COLOMBIANO OBLIGATORIO:
- Cuando sea pertinente, menciona entidades reales de Colombia:
  MinTIC, DANE, CRC, MinCiencias, SENA, DNP, Superintendencia de Industria y Comercio, etc.
- Usa cifras REALES y verificables de fuentes oficiales (DANE, MinTIC, CRC).
  Si no tienes un dato exacto confirmado, usa rangos o expresiones como
  "según estimaciones de..." o "alrededor del X%" en lugar de decimales precisos inventados.
  NUNCA inventes cifras exactas con decimales (ej: 23,7%) si no están en una fuente real.
- Menciona empresas, sectores o casos reales del país.
- Aterriza los conceptos globales al contexto nacional colombiano.
"""

_ESTILO_NATURAL = """
ESTILO DE REDACCIÓN:
- Redacción académica pero con voz propia: mezcla de análisis técnico y reflexión del autor.
- Evita frases vacías, genéricas o que suenen completamente automatizadas.
- Usa variedad de estructuras sintácticas; no comiences todos los párrafos igual.
- Incluye interpretaciones o valoraciones propias del autor en las secciones analíticas.
"""

def build_prompt(seccion, tema, info, tipo, norma, nivel, refs_manuales='', modo='rapido', objetivos_texto=''):
    instruccion_norma = NORMAS_INSTRUCCIONES.get(norma, NORMAS_INSTRUCCIONES['APA 7'])
    instruccion_tipo  = TIPOS_INSTRUCCIONES.get(tipo, TIPOS_INSTRUCCIONES['academico'])

    _NIVELES_INSTRUCCION = {
        'colegio':       "NIVEL COLEGIO/SECUNDARIA: Usa lenguaje claro y accesible. Oraciones cortas. Explica los conceptos desde cero sin asumir conocimiento previo. Evita tecnicismos; cuando los uses, defínelos. Extensión moderada.",
        'tecnico':       "NIVEL TÉCNICO/TECNOLÓGICO: Lenguaje práctico y orientado a la aplicación. Relaciona la teoría con ejemplos concretos del campo técnico. Incluye procedimientos y estándares cuando aplique. Evita abstracción excesiva.",
        'universitario': "NIVEL UNIVERSITARIO: Lenguaje académico formal con análisis crítico. Usa terminología disciplinar correctamente. Argumenta con citas y evidencias. Integra perspectivas teóricas y prácticas.",
        'posgrado':      "NIVEL POSGRADO/MAESTRÍA/DOCTORADO: Lenguaje altamente especializado. Profundidad teórica máxima. Discute debates académicos, limitaciones epistemológicas y brechas en la literatura. Cita fuentes primarias recientes (2022-2025). Análisis crítico riguroso.",
    }
    instruccion_nivel = _NIVELES_INSTRUCCION.get(nivel, _NIVELES_INSTRUCCION['universitario'])

    _MODOS_INSTRUCCION = {
        'rapido':     "",
        'automatico': f"\nINFORMACIÓN DEL AUTOR (apuntes/notas): Usa esta información como base y complementa con tu conocimiento académico:\n{info}\n",
        'manual':     f"\nINFORMACIÓN DEL AUTOR (texto propio para convertir): Toma este texto como el contenido central de la sección. Tu tarea es restructurarlo, formalizarlo y enriquecerlo con citas y lenguaje académico, SIN inventar hechos nuevos ni contradecir al autor:\n{info}\n",
    }
    instruccion_modo = _MODOS_INSTRUCCION.get(modo, "")

    # FIX 5: Detect subject area for domain-specific depth
    _AREAS_TEMATICAS = {
        'tecnologia': ['inteligencia artificial', 'ia ', ' ia,', 'machine learning', 'software', 'programacion',
                       'digital', 'internet', 'ciberseguridad', 'blockchain', 'datos', 'tic', 'algoritmo',
                       'automatizacion', 'robotica', 'computacion', 'redes', 'cloud', 'nube'],
        'salud':      ['salud', 'medicina', 'enfermedad', 'clinico', 'hospital', 'farmaco', 'vacuna',
                       'epidemia', 'pandemia', 'covid', 'paciente', 'quirurgico', 'terapia', 'diagnostico',
                       'nutricion', 'mental', 'psicologia', 'biologico', 'genetico', 'celula'],
        'educacion':  ['educacion', 'aprendizaje', 'enseñanza', 'pedagogia', 'curriculo', 'docente',
                       'universidad', 'colegio', 'estudiante', 'escuela', 'formacion', 'competencia',
                       'didactica', 'evaluacion', 'inclusion', 'desercion', 'alfabetizacion'],
        'medio_ambiente': ['ambiente', 'ecologia', 'sostenibilidad', 'cambio climatico', 'contaminacion',
                           'biodiversidad', 'residuos', 'energia renovable', 'deforestacion', 'agua',
                           'emisiones', 'carbono', 'mineria', 'petroleo', 'recurso natural'],
        'economia':   ['economia', 'finanzas', 'mercado', 'empresa', 'emprendimiento', 'inflacion',
                       'pib', 'comercio', 'exportacion', 'inversion', 'fiscal', 'tributario', 'banco',
                       'creditio', 'desempleo', 'laboral', 'productividad', 'competitividad'],
        'ciencias':   ['quimica', 'fisica', 'biologia', 'laboratorio', 'experimento', 'reaccion',
                       'molecula', 'atomo', 'celular', 'organico', 'inorganico', 'termodinamica',
                       'electromagnetismo', 'genetica', 'evolucion', 'taxonomia'],
        'derecho':    ['derecho', 'juridico', 'ley', 'norma', 'constitucion', 'penal', 'civil',
                       'contrato', 'litigio', 'jurisprudencia', 'tribunal', 'regulacion', 'politica publica'],
        'social':     ['sociedad', 'cultura', 'comunidad', 'migracion', 'genero', 'pobreza',
                       'desigualdad', 'violencia', 'paz', 'conflicto', 'derechos humanos', 'inclusion',
                       'diversidad', 'familia', 'poblacion'],
    }
    _AREA_INSTRUCCIONES = {
        'tecnologia':     "Área TECNOLOGÍA/IA: incluye arquitecturas, métricas técnicas (precisión, recall, F1), estándares ISO/IEEE aplicables, comparación de herramientas y casos de uso reales en Colombia (sector público y privado).",
        'salud':          "Área SALUD/MEDICINA: usa nomenclatura clínica correcta (CIE-10 si aplica), cifras epidemiológicas del INS o MinSalud, protocolos clínicos vigentes en Colombia, y referencias de revistas indexadas (PubMed, Scielo).",
        'educacion':      "Área EDUCACIÓN: referencia el sistema educativo colombiano (MEN, ICFES, SENA), incluye estadísticas de calidad educativa del DANE y MinEducación, y vincula con políticas como Ser Pilo Paga o modelos pedagógicos actuales.",
        'medio_ambiente': "Área MEDIO AMBIENTE: cita el IDEAM, MADS, CAR y acuerdos internacionales ratificados por Colombia (Acuerdo de París, COP28). Incluye datos de temperatura, deforestación o contaminación con fuente.",
        'economia':       "Área ECONOMÍA/FINANZAS: usa datos macroeconómicos del DANE, Banco de la República o DNP. Menciona sectores estratégicos colombianos (minería, agro, servicios, manufactura) con cifras del PIB sectoriales.",
        'ciencias':       "Área CIENCIAS NATURALES/LABORATORIO: usa nomenclatura IUPAC si aplica, describe procedimientos con precisión técnica, incluye ecuaciones o fórmulas cuando corresponda, y cita artículos de revistas científicas con DOI.",
        'derecho':        "Área DERECHO/JURÍDICO: cita la Constitución Política de Colombia de 1991, leyes, decretos y jurisprudencia real de la Corte Constitucional o Corte Suprema. Usa lenguaje jurídico técnico correcto.",
        'social':         "Área CIENCIAS SOCIALES: incluye perspectivas interseccionales cuando aplique, cita investigaciones cualitativas y cuantitativas, y contextualiza en la realidad sociopolítica colombiana (post-conflicto, desplazamiento, paz total).",
    }
    tema_lower = tema.lower()
    area_detectada = None
    max_matches = 0
    for area, keywords in _AREAS_TEMATICAS.items():
        matches = sum(1 for kw in keywords if kw in tema_lower)
        if matches > max_matches:
            max_matches = matches
            area_detectada = area
    instruccion_area = ("\n" + _AREA_INSTRUCCIONES[area_detectada] + "\n") if area_detectada else ""

    base = f"""Tipo de informe: {instruccion_tipo}
{instruccion_nivel}
Norma bibliográfica activa: {norma}.
{instruccion_norma}
{_CONTEXTO_COLOMBIA}
{_ESTILO_NATURAL}
{instruccion_area}{instruccion_modo}"""
    if seccion == 'introduccion':
        return base + f"""Escribe ÚNICAMENTE la INTRODUCCIÓN del informe sobre: "{tema}".
Estructura obligatoria (4-5 párrafos):
1. Contexto general del tema con datos estadísticos concretos y referencia a Colombia.
2. Justificación: por qué es relevante este tema en el panorama colombiano y latinoamericano.
3. Planteamiento del problema: ¿cuál es la brecha, tensión o necesidad que aborda el informe?
4. Alineación con los objetivos del informe: anuncia brevemente qué secciones siguen y qué aporta cada una.
- Al menos 2 citas en formato {norma}. Referencias entre 2019 y 2025.
- Menciona al menos una entidad colombiana relevante (MinTIC, DANE, CRC, etc.) si aplica al tema.
- NO incluyas el título de la sección, solo el contenido puro."""

    elif seccion == 'objetivos':
        return base + f"""Escribe ÚNICAMENTE los OBJETIVOS del informe sobre: "{tema}".
Usa este formato exacto:

OBJETIVO GENERAL:
[Un objetivo general claro, medible y redactado en infinitivo que capture la esencia del informe]

OBJETIVOS ESPECÍFICOS:
1. [Verbo en infinitivo + acción concreta + resultado esperado — directamente vinculado al tema]
2. [Verbo en infinitivo + acción concreta + resultado esperado]
3. [Verbo en infinitivo + acción concreta + resultado esperado]
4. [Verbo en infinitivo + acción concreta + resultado esperado]
5. [Verbo en infinitivo + acción concreta + resultado esperado]

Cada objetivo específico debe ser verificable y conectarse con una sección del informe (marco teórico, metodología, desarrollo, conclusiones o recomendaciones)."""

    elif seccion == 'marco_teorico':
        return base + f"""Escribe ÚNICAMENTE el MARCO TEÓRICO del informe sobre: "{tema}".
Estructura obligatoria (mínimo 5 párrafos):
1. Antecedentes históricos o evolución del tema.
2. Definiciones y conceptos clave con citas de autores reconocidos.
3. Teorías o modelos teóricos principales que sustentan el informe.
4. Estado del arte: investigaciones recientes (2020-2025) sobre el tema.
5. Contexto colombiano: cómo se ha estudiado o implementado este tema en Colombia.
- Al menos 5 citas en formato {norma} con años entre 2019 y 2025.
- Incluye autores latinoamericanos o colombianos cuando existan."""

    elif seccion == 'metodologia':
        # FIX 1: Metodología adaptada según tipo de informe real
        _METODO_POR_TIPO = {
            'laboratorio': f"""Este informe es un INFORME DE LABORATORIO sobre: "{tema}".
La metodología DEBE describir el procedimiento experimental real con estos 5 elementos:
1. Enfoque experimental: hipótesis planteada, variables independientes y dependientes, grupos de control.
2. Materiales y equipos: lista detallada de materiales usados (reactivos, instrumentos, software de medición).
3. Procedimiento paso a paso: descripción cronológica de los pasos del experimento con precisión técnica.
4. Técnicas de recolección de datos: cómo se midieron los resultados (tablas, gráficas, análisis estadístico).
5. Criterios de validez y limitaciones: posibles fuentes de error, condiciones del laboratorio, repetibilidad.
- Usa terminología técnica del área (química, biología, física, etc.) correctamente.
- Menciona al menos 2 protocolos o normas técnicas que apliquen al experimento.""",

            'pasantia': f"""Este informe es un INFORME DE PASANTÍA sobre: "{tema}".
La metodología describe el marco de trabajo de la práctica con estos 4 elementos:
1. Modalidad y duración: tipo de vinculación, empresa/entidad, horas totales, período.
2. Método de trabajo: metodologías usadas en la empresa (Scrum, Kanban, PMBOK, etc.) y cómo el pasante se integró.
3. Fuentes de información: documentos internos, manuales técnicos, capacitaciones recibidas, supervisión.
4. Evaluación y seguimiento: indicadores con los que se midió el desempeño del pasante, informes de avance.
- NO inventes trabajo de campo externo: la pasantía ocurrió dentro de la empresa.""",

            'proyecto': f"""Este informe es un INFORME DE PROYECTO sobre: "{tema}".
La metodología describe el marco de gestión del proyecto con estos 5 elementos:
1. Marco metodológico de gestión: enfoque usado (PMI/PMBOK, PRINCE2, metodología ágil, etc.) con justificación.
2. Fases del proyecto y cronograma: etapas definidas, hitos y entregables asociados a cada fase.
3. Técnicas de recolección y análisis: cómo se recopiló información para tomar decisiones (reuniones, informes, KPIs).
4. Gestión de riesgos: metodología para identificar y mitigar riesgos durante la ejecución.
5. Criterios de éxito y control de calidad: indicadores de desempeño definidos al inicio del proyecto.""",

            'ejecutivo': f"""Este informe es un INFORME EJECUTIVO sobre: "{tema}".
La metodología es concisa y orientada a la toma de decisiones con estos 3 elementos:
1. Fuentes consultadas: bases de datos empresariales, reportes del sector, benchmarks y analistas de referencia.
2. Criterios de análisis: variables clave evaluadas, período temporal, alcance geográfico (Colombia / Latam).
3. Limitaciones: acceso a información confidencial, oportunidad de los datos disponibles.""",

            'tesis': f"""Este informe tiene estructura de TESIS/MONOGRAFÍA sobre: "{tema}".
La metodología es el núcleo del trabajo académico y DEBE incluir estos 6 elementos:
1. Paradigma investigativo: positivista, interpretativo o socio-crítico — con fundamentación teórica.
2. Enfoque: cuantitativo, cualitativo o mixto — justificado epistemológicamente.
3. Tipo y alcance: exploratorio, descriptivo, correlacional o explicativo.
4. Población, muestra y muestreo: si aplica revisión sistemática, describe el protocolo PRISMA.
5. Instrumentos y técnicas: encuestas (con validación), entrevistas semiestructuradas, análisis documental, etc.
6. Procedimiento de análisis: software estadístico, análisis de contenido, triangulación, etc.""",
        }

        # Default para académico, investigativo y otros
        _metodo_default = f"""Este informe es de REVISIÓN DOCUMENTAL Y ANÁLISIS sobre: "{tema}".
REGLA DE CREDIBILIDAD: Este informe NO incluye investigación de campo real.
Basa la metodología en análisis de fuentes secundarias — NO inventes encuestas ni muestras.
Estructura obligatoria (4 párrafos):
1. Enfoque cualitativo-documental o descriptivo con justificación académica.
2. Tipo y alcance: investigación descriptiva basada en análisis de fuentes secundarias (informes
   del DANE, MinTIC, CRC, artículos de revistas indexadas, documentos de política pública).
3. Proceso de selección y análisis de fuentes: criterios de inclusión/exclusión, bases de datos
   consultadas (Scopus, Redalyc, Scielo, Google Scholar), período temporal cubierto.
4. Limitaciones honestas: acceso a datos primarios, actualización de fuentes, sesgo de publicación."""

        metodo_especifico = _METODO_POR_TIPO.get(tipo, _metodo_default)
        return base + f"""Escribe ÚNICAMENTE la METODOLOGÍA del informe sobre: "{tema}".
Adecua la complejidad al nivel educativo indicado.

{metodo_especifico}

- Al menos 2 citas en formato {norma} sobre metodología de investigación o del área específica.
- Tono académico, honesto y específico. NUNCA menciones datos de muestra inventados."""

    elif seccion == 'desarrollo':
        # FIX 2: Inject the actual generated objectives so the development directly addresses them
        bloque_objetivos = ""
        if objetivos_texto and objetivos_texto.strip():
            bloque_objetivos = f"""
OBJETIVOS DEL INFORME (YA GENERADOS — ÚSALOS COMO GUÍA OBLIGATORIA):
{objetivos_texto.strip()}

REGLA CRÍTICA: El desarrollo DEBE responder explícitamente a cada objetivo específico listado arriba.
Dedica al menos un párrafo sustancial a demostrar cómo se cumplió cada objetivo específico.
Usa frases de conexión directa como: "En relación con el objetivo específico N, que planteaba [objetivo]..."
NO puedes concluir el desarrollo sin haber abordado todos los objetivos específicos.
"""
        return base + f"""Escribe ÚNICAMENTE el DESARROLLO del informe de tipo "{tipo}" sobre: "{tema}".
{bloque_objetivos}
Estructura obligatoria (mínimo 7 párrafos):
1. Presentación del contexto y resultados principales con datos concretos y citas en formato {norma}.
2. Respuesta al Objetivo Específico 1: [desarrollar el hallazgo central con evidencia y citas].
3. Respuesta al Objetivo Específico 2: [segundo hallazgo con análisis comparativo].
4. Respuesta al Objetivo Específico 3: [tercer hallazgo con perspectiva crítica].
5. Análisis del contexto colombiano: aplica los resultados a la realidad de Colombia.
   - Cita datos del DANE, MinTIC, CRC u otras entidades institucionales relevantes.
   - Nombra empresas, sectores o casos reales del país con sus respectivas fuentes.
6. Tabla de datos estructurada — FORMATO OBLIGATORIO (usa EXACTAMENTE ##TABLE##/##ENDTABLE##):
   ##TABLE##
   TITULO: [Nombre descriptivo de la tabla]
   CABECERAS: Categoría | Valor | Año | Fuente
   FILA: [dato] | [dato] | [año] | [fuente]
   FILA: [dato] | [dato] | [año] | [fuente]
   FILA: [dato] | [dato] | [año] | [fuente]
   FILA: [dato] | [dato] | [año] | [fuente]
   ##ENDTABLE##
7. Análisis crítico del autor: ¿los hallazgos confirman la teoría?, ¿qué limitaciones existen?
- Al menos 5 citas en formato {norma} distribuidas en el texto. Fuentes entre 2021 y 2025.
- TODA cifra o estadística debe ir acompañada de su cita inmediatamente."""

    elif seccion == 'conclusiones':
        return base + f"""Escribe ÚNICAMENTE las CONCLUSIONES del informe sobre: "{tema}".
Usa este formato (5 conclusiones numeradas):
1. [Responde directamente al objetivo general — mínimo 3-4 oraciones explicando si se cumplió y por qué]
2. [Hallazgo más importante del desarrollo: dato clave o resultado central]
3. [Implicaciones prácticas para Colombia: qué debería cambiar o mejorar en el país]
4. [Limitaciones del estudio: qué no pudo resolverse y por qué]
5. [Perspectivas futuras: líneas de investigación o acciones recomendadas a futuro]
- Incluye al menos 1 opinión o interpretación propia del autor en las conclusiones.
- Usa un tono reflexivo y personal, no únicamente descriptivo."""

    elif seccion == 'recomendaciones':
        return base + f"""Escribe ÚNICAMENTE las RECOMENDACIONES del informe de tipo "{tipo}" sobre: "{tema}".
Usa este formato (5 recomendaciones numeradas):
1. [Dirigida al Ministerio o entidad pública colombiana relevante: acción concreta + justificación — 3 oraciones]
2. [Dirigida a empresas del sector privado en Colombia: acción específica + beneficio esperado]
3. [Dirigida a instituciones educativas o de investigación: formación, programas o estudios sugeridos]
4. [Dirigida a profesionales o practicantes del área: cambio de práctica o habilidad a desarrollar]
5. [Para futuras investigaciones: pregunta de investigación abierta o metodología a explorar]
- Cada recomendación debe ser accionable, dirigida a un actor específico y justificada."""

    # ── Secciones específicas para PASANTÍA ──────────────────────
    elif tipo == 'pasantia' and seccion == 'desarrollo':
        return base + f"""Escribe ÚNICAMENTE el DESARROLLO del INFORME DE PASANTÍA sobre: "{tema}".
Estructura obligatoria (7 párrafos mínimo):
1. Descripción de la empresa/organización: nombre, sector, misión, ubicación, tamaño y contexto en Colombia.
2. Rol y funciones del pasante: cargo, dependencia, supervisor y responsabilidades asignadas.
3. Actividades realizadas semana a semana o por etapas: descripción detallada de tareas ejecutadas.
4. Proyectos específicos en los que participó: objetivos del proyecto, contribución del pasante, resultado.
5. Herramientas, tecnologías y metodologías utilizadas durante la pasantía.
6. Competencias desarrolladas: habilidades técnicas, blandas y profesionales adquiridas.
7. Análisis crítico: ¿qué funcionó bien?, ¿qué desafíos enfrentó?, ¿qué aprendió que no está en los libros?
- Incluye al menos UNA tabla con cronograma de actividades o resumen de logros en formato ##TABLE##.
- Al menos 3 citas en formato {norma} sobre el área profesional o sector de la empresa."""""

    elif tipo == 'pasantia' and seccion == 'recomendaciones':
        return base + f"""Escribe ÚNICAMENTE las RECOMENDACIONES del INFORME DE PASANTÍA sobre: "{tema}".
Usa este formato (5 recomendaciones numeradas, con destinatario explícito):
1. [Para la empresa/organización donde se realizó la pasantía: mejora de proceso, área o práctica concreta que el pasante identificó]
2. [Para futuros pasantes que vayan a la misma área: preparación, actitud o conocimientos recomendados]
3. [Para la institución educativa: ajustes al pensum o preparación previa que facilite la práctica]
4. [Para el área o dependencia específica: cambio operativo o de gestión que mejoraría el desempeño del equipo]
5. [Para el propio autor: plan de desarrollo profesional personal a partir de esta experiencia]
- Cada recomendación debe ser específica, accionable y basada en la experiencia descrita."""""

    elif tipo == 'proyecto' and seccion == 'desarrollo':
        return base + f"""Escribe ÚNICAMENTE el DESARROLLO del INFORME DE PROYECTO sobre: "{tema}".
Estructura obligatoria (7 párrafos mínimo):
1. Descripción general del proyecto: alcance, justificación y contexto organizacional o académico.
2. Fases o etapas del proyecto con fechas clave: planificación, ejecución, seguimiento, cierre.
3. Recursos utilizados: humanos (roles del equipo), tecnológicos, financieros y materiales.
4. Entregables principales: qué se produjo en cada etapa y cuál fue su estado de completitud.
5. Gestión de riesgos: riesgos identificados, probabilidad, impacto y medidas de mitigación aplicadas.
6. Estado de avance y resultados: ¿se cumplieron los hitos?, indicadores de éxito o KPIs del proyecto.
7. Lecciones aprendidas y análisis crítico: desviaciones del plan original y cómo se gestionaron.
- Incluye al menos UNA tabla en formato ##TABLE## con hitos, fechas y estado (completado/en progreso/pendiente).
- Al menos 3 citas en formato {norma} sobre gestión de proyectos o el área del proyecto."""""

    elif tipo == 'proyecto' and seccion == 'recomendaciones':
        return base + f"""Escribe ÚNICAMENTE las RECOMENDACIONES del INFORME DE PROYECTO sobre: "{tema}".
Usa este formato (5 recomendaciones numeradas):
1. [Para el equipo del proyecto: mejora de proceso, comunicación o metodología en futuros proyectos similares]
2. [Para la dirección o patrocinador: decisión estratégica sobre continuación, escalamiento o replicación]
3. [Para la gestión de riesgos: riesgos no contemplados que deberían incluirse en proyectos similares]
4. [Para la gestión de recursos: optimización de tiempos, costos o asignación de roles]
5. [Para futuras investigaciones o proyectos relacionados: líneas de acción o preguntas abiertas]
- Cada recomendación debe ser específica, accionable y vinculada a los hallazgos del proyecto."""""

    elif seccion == 'referencias':
        refs_extra = f"\nIncluye o adapta estas referencias del autor:\n{refs_manuales}" if refs_manuales else ""

        # Instrucciones de formato específicas por norma
        _FORMATO_REFS = {
            'APA 7': """FORMATO APA 7 ESTRICTO:
   - Entidades: Sigla. (año). Título del documento. Nombre completo de la entidad. URL
   - Personas: Apellido, I. I. (año). Título. Revista, vol(num), págs. https://doi.org/...
   - Orden: ALFABÉTICO por primer apellido o sigla de entidad
   - Sin numeración. Sin negrita. DOI en formato https://doi.org/xxxxx""",
            'APA 6': """FORMATO APA 6 ESTRICTO:
   - Entidades: Nombre Entidad. (año). Título del documento. Ciudad: Editorial. Recuperado de URL
   - Personas: Apellido, I. I. (año). Título. Revista, vol(num), págs.
   - Orden: ALFABÉTICO. Sin numeración. Sin negrita.""",
            'ICONTEC': """FORMATO ICONTEC (NTC 5613) ESTRICTO:
   - Entidades: NOMBRE ENTIDAD. Título del documento. Ciudad: Editorial, año. Disponible en: URL
   - Personas: APELLIDO, Nombre. Título en cursiva. Ed. Ciudad: Editorial, año.
   - Orden: ORDEN DE APARICIÓN en el texto (numeradas 1, 2, 3...)
   - Separador entre autores: punto y coma (;)""",
            'IEEE': """FORMATO IEEE ESTRICTO:
   - Artículos: [N] I. Apellido, "Título," Nombre Revista, vol. X, no. X, pp. XX-XX, año. doi: ...
   - Libros: [N] I. Apellido, Título del libro. Ciudad: Editorial, año.
   - Entidades: [N] Nombre Entidad, "Título del documento," año. [Online]. Available: URL
   - Orden: NUMÉRICO por orden de aparición en el texto""",
            'Vancouver': """FORMATO VANCOUVER ESTRICTO:
   - Artículos: N. Apellido AB, Apellido CD. Título del artículo. Abrev Revista. año;vol(num):págs.
   - Libros: N. Apellido AB. Título. Edición. Ciudad: Editorial; año.
   - Entidades: N. Nombre Entidad. Título [Internet]. Ciudad: Entidad; año [citado año mes día]. Disponible en: URL
   - Orden: NUMÉRICO por orden de aparición. Máximo 6 autores, luego "et al."
   - Sin negrita. Sin DOI en formato URL, usar: doi:xxxxxxxxx""",
            'Chicago': """FORMATO CHICAGO 17ª EDICIÓN ESTRICTO:
   - Artículos: Apellido, Nombre. año. "Título del artículo." Nombre Revista vol, no. num: págs. URL/DOI.
   - Libros: Apellido, Nombre. año. Título en cursiva. Ciudad: Editorial.
   - Entidades: Nombre Entidad. año. "Título del documento." URL.
   - Orden: ALFABÉTICO por apellido""",
            'MLA': """FORMATO MLA 9ª EDICIÓN ESTRICTO:
   - Artículos: Apellido, Nombre. "Título del artículo." Nombre Revista, vol. X, no. X, año, pp. XX-XX. DOI/URL.
   - Libros: Apellido, Nombre. Título en cursiva. Editorial, año.
   - Entidades: Nombre Entidad. "Título del documento." Año, URL.
   - Sección: Works Cited (no "Referencias")
   - Orden: ALFABÉTICO""",
            'Harvard': """FORMATO HARVARD ESTRICTO:
   - Artículos: Apellido, I. (año) 'Título del artículo', Nombre Revista, vol. X, no. X, pp. XX-XX.
   - Libros: Apellido, I. (año) Título en cursiva. Ciudad: Editorial.
   - Entidades: Nombre Entidad (año) Título del documento. Ciudad: Entidad. Available at: URL (Accessed: fecha).
   - Orden: ALFABÉTICO por apellido""",
        }
        formato_norma = _FORMATO_REFS.get(norma, _FORMATO_REFS['APA 7'])

        # Ejemplos de entidades colombianas adaptados por norma
        _EJEMPLOS_COLOMBIA = {
            'APA 7':    "DANE. (2023). Nombre del informe relevante al tema. Departamento Administrativo Nacional de Estadística. https://www.dane.gov.co",
            'APA 6':    "DANE. (2023). Nombre del informe relevante al tema. Bogotá: DANE. Recuperado de https://www.dane.gov.co",
            'ICONTEC':  "DEPARTAMENTO ADMINISTRATIVO NACIONAL DE ESTADÍSTICA (DANE). Nombre del informe relevante al tema. Bogotá: DANE, 2023. Disponible en: https://www.dane.gov.co",
            'IEEE':     '[N] DANE, "Nombre del informe relevante al tema," 2023. [Online]. Available: https://www.dane.gov.co',
            'Vancouver':"N. DANE. Nombre del informe [Internet]. Bogotá: DANE; 2023. Disponible en: https://www.dane.gov.co",
            'Chicago':  "DANE. 2023. \"Nombre del informe relevante al tema.\" https://www.dane.gov.co.",
            'MLA':      'DANE. "Nombre del informe relevante al tema." 2023, https://www.dane.gov.co.',
            'Harvard':  "DANE (2023) Nombre del informe relevante al tema. Bogotá: DANE. Available at: https://www.dane.gov.co",
        }
        ejemplo_colombia = _EJEMPLOS_COLOMBIA.get(norma, _EJEMPLOS_COLOMBIA['APA 7'])

        return base + f"""Genera ÚNICAMENTE la lista de REFERENCIAS BIBLIOGRÁFICAS en norma {norma} sobre: "{tema}".
{refs_extra}
CRITERIOS OBLIGATORIOS — cumple TODOS sin excepción:

1. CANTIDAD Y FECHA: Entre 10 y 12 referencias. TODAS de años 2021-2025. Ninguna anterior a 2021.

2. COHERENCIA: Las referencias deben corresponder EXACTAMENTE a las fuentes citadas en el informe.
   Si el texto menciona DANE, MinTIC, CRC, SENA, DNP, MinCiencias u otras entidades colombianas,
   su referencia DEBE aparecer en la lista.

3. ENTIDADES COLOMBIANAS (mínimo 3): Adapta el formato a la norma {norma}.
   Ejemplo de cómo citar el DANE en norma {norma}:
   {ejemplo_colombia}
   Haz lo mismo con MinTIC, CRC, DNP u otras entidades pertinentes al tema "{tema}".

4. ARTÍCULOS ACADÉMICOS (mínimo 4): con revista, volumen, número, páginas y DOI real.
   Deben ser directamente sobre el tema del informe.

5. ORGANISMO INTERNACIONAL (mínimo 1): CEPAL, BID, UNESCO, OCDE u otro, con URL oficial.
   Formatea también este organismo en norma {norma}.

6. {formato_norma}

7. NO incluyas referencias de temas NO relacionados con el informe.
   NO mezcles formatos de otras normas. Aplica SOLO {norma} en todas las referencias.

Sin título de sección, sin preámbulo, sin explicación. Solo las referencias."""

    # ── Tipos especiales: Pasantía ──────────────────────────────
    elif tipo == 'pasantia' and seccion in ('introduccion', 'objetivos', 'marco_teorico',
                                             'metodologia', 'desarrollo', 'conclusiones',
                                             'recomendaciones'):
        # Para pasantía, si no hay un prompt de sección específico arriba, usar estructura de pasantía
        pass  # los prompts anteriores ya manejan las secciones, la diferencia viene de instruccion_tipo

    # ── Tipos especiales: Proyecto ──────────────────────────────
    elif tipo == 'proyecto' and seccion in ('introduccion', 'objetivos', 'marco_teorico',
                                             'metodologia', 'desarrollo', 'conclusiones',
                                             'recomendaciones'):
        pass  # ídem

    return ""

# ============================================================
# GENERAR SECCIÓN
# ============================================================
def generar_seccion(seccion, tema, info_extra, tipo_informe, norma, nivel, refs_manuales='', modo='rapido', objetivos_texto='', contexto_refs=''):
    prompt = build_prompt(seccion, tema, info_extra, tipo_informe, norma, nivel, refs_manuales, modo, objetivos_texto)
    if not prompt:
        return None

    _nivel_sistema = {
        'colegio':       "para un estudiante de colegio/secundaria: usa lenguaje simple, claro y accesible, sin tecnicismos innecesarios",
        'tecnico':       "para un estudiante técnico/tecnológico: lenguaje práctico, orientado a la aplicación profesional",
        'universitario': "para un estudiante universitario: lenguaje académico formal con análisis crítico y terminología disciplinar",
        'posgrado':      "para un estudiante de posgrado/maestría/doctorado: máxima profundidad teórica, debate académico y análisis epistemológico",
    }.get(nivel, "para un estudiante universitario")

    _modo_sistema = {
        'rapido':     "",
        'automatico': " El autor te proporcionó sus apuntes o notas; úsalos como base y complementa con conocimiento académico.",
        'manual':     " El autor te proporcionó su propio texto; tu tarea es formalizarlo y enriquecerlo con lenguaje académico y citas, sin inventar hechos nuevos.",
    }.get(modo, "")

    # Bloque de fuentes reales — va al INICIO del prompt de usuario (más peso que system)
    _prefijo_refs = ""
    if contexto_refs and contexto_refs.strip() and seccion not in ('referencias',):
        _apellidos_lista = re.findall(r'\[\d+\]\s+([A-ZÁÉÍÓÚÑa-záéíóúñ\-]+)', contexto_refs)
        _apellidos_str   = ", ".join(_apellidos_lista[:12]) if _apellidos_lista else "los listados arriba"
        _prefijo_refs = (
            f"FUENTES VERIFICADAS — LISTA COMPLETA Y ÚNICA PERMITIDA:\n"
            f"{contexto_refs}\n"
            f"═══════════════════════════════════════════════════════\n"
            f"REGLAS ABSOLUTAS DE CITACIÓN:\n"
            f"1. SOLO puedes citar autores de la lista de arriba. Apellidos permitidos: {_apellidos_str}.\n"
            f"2. PROHIBIDO ABSOLUTO inventar autores: ni García, ni López, ni Martínez ni NINGÚN apellido fuera de la lista.\n"
            f"3. Si necesitas más citas, reutiliza el mismo autor con distintos puntos.\n"
            f"4. Entidades colombianas (DANE, MinSalud, MinTIC, MinCiencias, OPS) solo para datos estadísticos generales.\n"
            f"5. Cada cita en el texto DEBE corresponder a un ítem de la lista de arriba.\n"
            f"═══════════════════════════════════════════════════════\n\n"
        )

    system_prompt = (
        f"Eres un experto en redacción académica en español, especializado en norma {norma} "
        f"y profundo conocedor del contexto colombiano y latinoamericano. "
        f"Escribes {_nivel_sistema}.{_modo_sistema} "
        f"REGLA CRÍTICA DE CITAS: TODA afirmación factual, estadística o concepto teórico "
        f"DEBE tener una cita en el texto en formato {norma}. "
        f"SOLO puedes citar autores que estén en la lista de FUENTES VERIFICADAS que recibirás. "
        f"No puedes escribir un dato o afirmación sin su respectiva cita inmediatamente después. "
        f"Las secciones de introducción, marco teórico y desarrollo deben tener al menos "
        f"una cita por párrafo, usando exclusivamente los autores de la lista verificada. "
        "REGLA CRÍTICA DE VERACIDAD: NUNCA inventes datos, cifras, estadísticas ni nombres de estudios. "
        "Si no tienes el dato exacto de una fuente real, escribe: "
        "'según estimaciones recientes' o 'de acuerdo con reportes del sector (año aproximado)'. "
        "NUNCA uses porcentajes con decimales inventados (ej: 23,7% o 41,3%) — usa rangos: 'entre el 20 y 25%'. "
        "Si citas a MinTIC, DANE, CRC u otra entidad colombiana, el dato debe ser real y verificable. "
        "Cuando aplique al tema, incluyes datos de Colombia citando fuentes como MinTIC, DANE, CRC o MinCiencias. "
        "Usas referencias de años 2021-2025. "
        f"Para referencias en norma {norma}: aplica formato estrictamente correcto. "
        "Balanceas el análisis técnico con interpretaciones propias del autor. "
        "Respondes SOLO con el contenido solicitado, sin títulos de sección, sin preámbulos, sin asteriscos (**), sin markdown."
    )

    # Inyectar el prefijo de fuentes reales AL INICIO del prompt de usuario
    prompt = _prefijo_refs + prompt

    contenido = llamar_deepseek(prompt, system_prompt=system_prompt, max_tokens=3000)
    if not contenido:
        return None
    contenido = contenido.strip()

    # ── Validar autores reales y detectar inventados ──
    if contexto_refs and contexto_refs.strip() and seccion not in ('referencias',):
        apellidos_reales = re.findall(r'\[\d+\]\s+([A-ZÁÉÍÓÚÑa-záéíóúñ\-]+)', contexto_refs)
        if apellidos_reales:
            apellidos_usados = sum(1 for ap in apellidos_reales if ap.lower() in contenido.lower())
            citas_en_texto   = re.findall(r'\(([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+)[^)]{0,40}\d{4}\)', contenido)
            entidades_ok     = {'dane', 'minsalud', 'mintic', 'minciencias', 'ops', 'oms',
                                'who', 'fedegan', 'crc', 'colciencias', 'sena'}
            citas_inventadas = [c for c in citas_en_texto
                                if not any(ap.lower() in c.lower() for ap in apellidos_reales)
                                and c.lower().split()[0] not in entidades_ok]
            if apellidos_usados == 0 or len(citas_inventadas) > 2:
                razon = "no usó autores reales" if apellidos_usados == 0 else f"inventó autores: {citas_inventadas[:3]}"
                logger.warning(f"Seccion '{seccion}': {razon}. Reintentando...")
                prompt_retry = (
                    f"FUENTES VERIFICADAS — ÚNICAS PERMITIDAS:\n{contexto_refs}\n\n"
                    f"PROBLEMA: {razon}.\n"
                    f"Apellidos permitidos: {chr(10).join(apellidos_reales[:10])}.\n"
                    f"Reescribe usando SOLO esos autores. Para datos sin autor usa (DANE, año) o (MinSalud, año).\n\n"
                    f"SECCIÓN:\n"
                    f"{prompt.split('═══')[-1] if '═══' in prompt else prompt.split('━━━')[-1] if '━━━' in prompt else prompt}"
                )
                contenido2 = llamar_deepseek(prompt_retry, system_prompt=system_prompt, max_tokens=3000)
                if contenido2:
                    ap2  = sum(1 for ap in apellidos_reales if ap.lower() in contenido2.lower())
                    inv2 = [c for c in re.findall(r'\(([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+)[^)]{0,40}\d{4}\)', contenido2)
                            if not any(ap.lower() in c.lower() for ap in apellidos_reales)
                            and c.lower().split()[0] not in entidades_ok]
                    if ap2 >= apellidos_usados and len(inv2) <= len(citas_inventadas):
                        contenido = contenido2.strip()
                        logger.info(f"Reintento OK: {ap2} reales, {len(inv2)} inventados")

    # Validar citas — 1 reintento si faltan
    val = validar_citas(contenido, seccion, norma)
    if not val['ok']:
        logger.warning(f"Seccion '{seccion}': {val['citas_encontradas']}/{val['minimo']} citas. Reintentando...")
        p2 = prompt + (
            f"\n\nATENCION: el texto anterior tiene solo {val['citas_encontradas']} citas en formato {norma}. "
            f"Necesitas minimo {val['minimo']}. Reescribe incluyendo al menos una cita por parrafo."
        )
        c2 = llamar_deepseek(p2, system_prompt=system_prompt, max_tokens=3000)
        if c2:
            v2 = validar_citas(c2.strip(), seccion, norma)
            if v2['citas_encontradas'] > val['citas_encontradas']:
                contenido = c2.strip()
                logger.info(f"Reintento mejoro citas: {v2['citas_encontradas']}")

    # Validar tabla en desarrollo — 1 reintento si falta
    if seccion == 'desarrollo' and not validar_tabla_en_desarrollo(contenido):
        logger.warning("Desarrollo sin tabla. Reintentando...")
        p3 = prompt + (
            "\n\nATENCION: falta la tabla obligatoria ##TABLE##...##ENDTABLE##. "
            "Reescribe el punto 6 incluyendo la tabla exactamente en ese formato."
        )
        c3 = llamar_deepseek(p3, system_prompt=system_prompt, max_tokens=3200)
        if c3 and validar_tabla_en_desarrollo(c3.strip()):
            contenido = c3.strip()
            logger.info("Reintento incluyo tabla")

    logger.info(f"Seccion '{seccion}' generada: {len(contenido)} chars")
    return contenido

def _construir_contexto_refs(refs: list, norma: str) -> str:
    """
    Convierte la lista de papers reales en un bloque de texto que DeepSeek
    puede usar como fuentes al redactar. Incluye autor, año, título y DOI
    para que las citas en el texto sean verificables.
    """
    if not refs:
        return ""
    lineas = ["FUENTES REALES VERIFICADAS (úsalas para las citas del texto):"]
    for i, r in enumerate(refs, 1):
        autores = r.get("autores", [])
        if autores:
            apellidos = ", ".join(a["apellido"] for a in autores[:3])
            if len(autores) > 3:
                apellidos += " et al."
        else:
            apellidos = "Autor desconocido"
        anio   = r.get("anio", "s.f.")
        titulo = r.get("titulo", "")[:120]
        doi    = r.get("doi", "")
        revista = r.get("revista", "") or r.get("editorial", "")
        linea = f"[{i}] {apellidos} ({anio}). {titulo}."
        if revista:
            linea += f" {revista}."
        if doi:
            linea += f" DOI: {doi}"
        lineas.append(linea)
    lineas.append(
        "\nREGLA CRÍTICA: Cita SOLO los autores listados arriba. "
        "Usa su apellido y año exactos (ej: García, 2023). "
        "NO inventes autores, apellidos ni años que no aparezcan en esta lista."
    )
    return "\n".join(lineas)


def generar_informe_completo(tema, info_extra, tipo_informe, norma, nivel, modo='rapido'):
    """
    Genera el informe con referencias reales de CrossRef/OpenAlex.
    Orden: 1) buscar papers reales → 2) inyectarlos en cada sección → 3) generar texto.
    Así las citas en el texto corresponden a fuentes reales verificadas.
    """
    claves = ['introduccion', 'objetivos', 'marco_teorico', 'metodologia',
              'desarrollo', 'conclusiones', 'recomendaciones', 'referencias']
    secciones = {c: '' for c in claves}

    # ── PASO 1: buscar papers reales ANTES de generar cualquier sección ──
    logger.info(f"Buscando referencias reales para '{tema[:50]}'...")
    refs_reales = []
    refs_texto  = None
    refs_total  = 0
    try:
        refs_reales = buscar_referencias_reales(tema, cantidad_total=12)
        if refs_reales:
            refs_texto = formatear_referencias(refs_reales, norma)
            refs_total = len(refs_reales)
            logger.info(f"Referencias reales obtenidas: {refs_total}")
    except Exception as e:
        logger.error(f"Error buscando referencias reales: {e}")

    # Construir bloque de contexto que se inyectará en cada prompt
    contexto_refs = _construir_contexto_refs(refs_reales, norma)

    # ── PASO 2: generar secciones con las fuentes reales como contexto ──
    claves_ia = [c for c in claves if c not in ('desarrollo', 'referencias')]

    def _generar(clave):
        resultado = generar_seccion(
            clave, tema, info_extra, tipo_informe, norma, nivel,
            modo=modo, contexto_refs=contexto_refs
        )
        return clave, resultado or ''

    with ThreadPoolExecutor(max_workers=5) as executor:
        futuros_ia = {executor.submit(_generar, c): c for c in claves_ia}
        for futuro in as_completed(futuros_ia):
            try:
                clave, contenido = futuro.result()
                secciones[clave] = contenido
                logger.info(f"Seccion '{clave}' completada ({len(contenido)} chars)")
            except Exception as e:
                clave = futuros_ia[futuro]
                logger.error(f"Error en seccion '{clave}': {e}")

    # Desarrollo con objetivos inyectados + contexto de fuentes reales
    objetivos_generados = secciones.get('objetivos', '')
    secciones['desarrollo'] = generar_seccion(
        'desarrollo', tema, info_extra, tipo_informe, norma, nivel,
        modo=modo, objetivos_texto=objetivos_generados, contexto_refs=contexto_refs
    ) or ''
    logger.info(f"Seccion 'desarrollo' completada ({len(secciones['desarrollo'])} chars)")

    # ── PASO 3: referencias — usar las reales si hay suficientes ──
    if refs_texto and refs_total >= 3:
        secciones['referencias'] = refs_texto
        logger.info(f"Referencias reales asignadas: {refs_total}")
    else:
        logger.warning(f"Referencias reales insuficientes ({refs_total}), usando IA con aviso")
        refs_ia = generar_seccion(
            'referencias', tema, info_extra, tipo_informe, norma, nivel,
            refs_manuales=info_extra, modo=modo
        ) or ''
        aviso = (
            "[NOTA: No se encontraron suficientes referencias verificadas en bases de datos "
            "académicas. Las siguientes referencias fueron generadas por IA y deben verificarse "
            "antes de su uso académico.]\n\n"
        )
        secciones['referencias'] = aviso + refs_ia

    # Revisión de coherencia final
    secciones = _revisar_coherencia(secciones, tema, norma)
    # CAMBIO 2: Devolver también las referencias reales
    return secciones, refs_reales


def _revisar_coherencia(secciones: dict, tema: str, norma: str) -> dict:
    """Verifica que conclusiones respondan a los objetivos; regenera si la cobertura es baja."""
    objetivos   = secciones.get('objetivos', '')
    conclusiones = secciones.get('conclusiones', '')
    desarrollo  = secciones.get('desarrollo', '')
    if not objetivos or not conclusiones:
        return secciones
    primera_linea = objetivos.strip().split('\n')[0].lower()
    palabras_clave = [w for w in primera_linea.split() if len(w) > 5]
    concl_lower    = conclusiones.lower()
    encontradas    = sum(1 for w in palabras_clave if w in concl_lower)
    cobertura      = encontradas / max(len(palabras_clave), 1)
    if cobertura < 0.25:
        logger.warning(f"Coherencia baja ({cobertura:.0%}). Regenerando conclusiones...")
        system_c = (
            f"Eres un revisor académico especializado en norma {norma}. "
            f"Reescribe las conclusiones para que respondan directamente a los objetivos. "
            f"No inventes datos nuevos."
        )
        prompt_c = (
            f'Informe sobre: "{tema}".\n\nOBJETIVOS:\n{objetivos}\n\n'
            f'DESARROLLO (resumen):\n{desarrollo[:800]}...\n\n'
            f'CONCLUSIONES ACTUALES (reescribir):\n{conclusiones}\n\n'
            f'Reescribe las conclusiones (5 puntos numerados) respondiendo a cada objetivo. '
            f'Sin título de sección.'
        )
        nuevas = llamar_deepseek(prompt_c, system_prompt=system_c, max_tokens=1500)
        if nuevas and len(nuevas.strip()) > 200:
            secciones['conclusiones'] = nuevas.strip()
            logger.info("Conclusiones regeneradas por baja coherencia")
    return secciones

# ============================================================
# GENERAR PDF
# ============================================================
def generar_pdf(datos_usuario, secciones):
    nombre      = datos_usuario.get('nombre', 'Estudiante')
    autores_extra = datos_usuario.get('autores_extra', [])
    tema        = datos_usuario.get('tema', 'Tema')
    asignatura  = datos_usuario.get('asignatura', '')
    profesor    = datos_usuario.get('profesor', '')
    institucion = datos_usuario.get('institucion', '')
    ciudad      = datos_usuario.get('ciudad', '')
    fecha       = datos_usuario.get('fecha', datetime.now().strftime('%d/%m/%Y'))
    norma       = datos_usuario.get('norma', 'APA 7')

    filename = f"informe_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}.pdf"
    filepath = os.path.join('informes_generados', filename)

    styles = getSampleStyleSheet()

    styles.add(ParagraphStyle(
        name='TextoJustificado',
        parent=styles['Normal'],
        alignment=TA_JUSTIFY,
        fontSize=11,
        fontName='Times-Roman',
        spaceAfter=10,
        leading=18,
        firstLineIndent=18
    ))
    styles.add(ParagraphStyle(
        name='Titulo1',
        parent=styles['Heading1'],
        fontSize=14,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#1a365d'),
        spaceBefore=20,
        spaceAfter=10,
        alignment=TA_LEFT
    ))
    styles.add(ParagraphStyle(
        name='TituloPortada',
        parent=styles['Normal'],
        fontSize=22,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#1a365d'),
        alignment=TA_CENTER,
        spaceAfter=12
    ))
    styles.add(ParagraphStyle(
        name='SubtituloPortada',
        parent=styles['Normal'],
        fontSize=14,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#2d4a7a'),
        alignment=TA_CENTER,
        spaceAfter=8
    ))
    styles.add(ParagraphStyle(
        name='TextoCentrado',
        parent=styles['Normal'],
        fontSize=11,
        fontName='Times-Roman',
        alignment=TA_CENTER,
        spaceAfter=5
    ))
    styles.add(ParagraphStyle(
        name='ObjetivoTitulo',
        parent=styles['Normal'],
        alignment=TA_LEFT,
        fontSize=11,
        fontName='Helvetica-Bold',
        spaceBefore=14,
        spaceAfter=4,
        textColor=colors.HexColor('#1a365d'),
        leading=16
    ))
    styles.add(ParagraphStyle(
        name='ObjetivoItem',
        parent=styles['Normal'],
        alignment=TA_JUSTIFY,
        fontSize=11,
        fontName='Times-Roman',
        spaceBefore=8,
        spaceAfter=8,
        leftIndent=16,
        leading=18
    ))
    styles.add(ParagraphStyle(
        name='TextoLista',
        parent=styles['Normal'],
        alignment=TA_LEFT,
        fontSize=11,
        fontName='Times-Roman',
        spaceAfter=6,
        leftIndent=20,
        leading=16
    ))
    styles.add(ParagraphStyle(
        name='Referencia',
        parent=styles['Normal'],
        alignment=TA_JUSTIFY,
        fontSize=10.5,
        fontName='Times-Roman',
        spaceAfter=8,
        leading=16,
        leftIndent=30,        # sangría francesa: cuerpo desplazado a la derecha
        firstLineIndent=-30   # primera línea al margen izquierdo
    ))

    doc = SimpleDocTemplate(
        filepath,
        pagesize=A4,
        rightMargin=2.5*cm,
        leftMargin=2.5*cm,
        topMargin=2.5*cm,
        bottomMargin=2.5*cm
    )
    story = []

    # PORTADA
    story.append(Spacer(1, 2.0*inch))
    story.append(Paragraph("INFORME ACADÉMICO", styles['TituloPortada']))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor('#1a365d'), spaceAfter=12))
    story.append(Paragraph(tema.upper(), styles['SubtituloPortada']))
    story.append(Spacer(1, 1.2*inch))
    story.append(Paragraph(f"<b>Presentado por:</b> {nombre}", styles['TextoCentrado']))
    for autor in autores_extra:
        if autor.get('nombre'):
            linea = autor['nombre']
            if autor.get('cargo'):
                linea += f" — {autor['cargo']}"
            story.append(Paragraph(linea, styles['TextoCentrado']))
    if asignatura:
        story.append(Paragraph(f"<b>Asignatura:</b> {asignatura}", styles['TextoCentrado']))
    if profesor:
        story.append(Paragraph(f"<b>Docente:</b> {profesor}", styles['TextoCentrado']))
    if institucion:
        story.append(Paragraph(f"<b>Institución:</b> {institucion}", styles['TextoCentrado']))
    if ciudad:
        story.append(Paragraph(f"<b>Ciudad:</b> {ciudad}", styles['TextoCentrado']))
    story.append(Paragraph(f"<b>Fecha:</b> {fecha}", styles['TextoCentrado']))
    story.append(Paragraph(f"<b>Norma bibliográfica:</b> {norma}", styles['TextoCentrado']))
    story.append(PageBreak())

    # ÍNDICE
    story.append(Paragraph("ÍNDICE", styles['Titulo1']))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#cccccc'), spaceAfter=8))
    for idx in ["1. Introducción", "2. Objetivos", "3. Marco Teórico",
                "4. Metodología", "5. Desarrollo", "6. Conclusiones",
                "7. Recomendaciones", "8. Referencias Bibliográficas"]:
        story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;{idx}", styles['TextoLista']))
    story.append(PageBreak())

    # SECCIONES
    secciones_orden = [
        ("1. INTRODUCCIÓN",               'introduccion'),
        ("2. OBJETIVOS",                  'objetivos'),
        ("3. MARCO TEÓRICO",              'marco_teorico'),
        ("4. METODOLOGÍA",                'metodologia'),
        ("5. DESARROLLO",                 'desarrollo'),
        ("6. CONCLUSIONES",               'conclusiones'),
        ("7. RECOMENDACIONES",            'recomendaciones'),
        ("8. REFERENCIAS BIBLIOGRÁFICAS", 'referencias'),
    ]

    for titulo, clave in secciones_orden:
        story.append(Paragraph(titulo, styles['Titulo1']))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e0e0e0'), spaceAfter=8))
        contenido_raw = secciones.get(clave, '')
        contenido = limpiar_para_pdf(contenido_raw)

        if contenido and len(contenido) > 50:
            # ── Renderizado especial para OBJETIVOS ─────────────────
            if clave == 'objetivos':
                lineas = [l.strip() for l in re.split(r'\n+', contenido) if l.strip()]
                for linea in lineas:
                    if re.match(r'^OBJETIVO\s+(GENERAL|ESPEC)', linea, re.IGNORECASE):
                        story.append(Paragraph(linea, styles['ObjetivoTitulo']))
                    elif re.match(r'^\d+\.', linea):
                        story.append(Paragraph(linea, styles['ObjetivoItem']))
                    else:
                        story.append(Paragraph(linea, styles['TextoJustificado']))
                story.append(PageBreak())
                continue

            contenido_proc, tablas = extraer_tablas(contenido)
            parrafos = re.split(r'\n{2,}', contenido_proc)
            if len(parrafos) == 1:
                parrafos = contenido_proc.split('\n')
            for p in parrafos:
                p = p.strip()
                if not p:
                    continue
                # Detectar marcador de tabla
                m_tabla = re.match(r'^__TABLA_(\d+)__$', p)
                if m_tabla and tablas:
                    idx_t = int(m_tabla.group(1))
                    if idx_t < len(tablas):
                        t = tablas[idx_t]
                        if t.get("titulo"):
                            story.append(Spacer(1, 6))
                            story.append(Paragraph(
                                f"<i>{t['titulo']}</i>",
                                ParagraphStyle('TituloTabla', parent=styles['Normal'],
                                    fontSize=9, fontName='Helvetica-Bold',
                                    textColor=colors.HexColor('#1a365d'),
                                    spaceAfter=4, spaceBefore=10)
                            ))
                        # Construir datos de tabla
                        cabeceras = t.get("cabeceras", [])
                        filas     = t.get("filas", [])
                        # Normalizar ancho de columnas
                        n_cols = max(len(cabeceras), max((len(f) for f in filas), default=0))
                        if n_cols == 0:
                            continue
                        def normalizar(fila, n):
                            return fila[:n] + [''] * max(0, n - len(fila))
                        data = [normalizar(cabeceras, n_cols)] + [normalizar(f, n_cols) for f in filas]
                        col_w = (15.5 / n_cols)  # cm, A4 ancho útil ~15.5cm
                        tbl = Table(data, colWidths=[col_w * rl_cm] * n_cols, repeatRows=1)
                        tbl.setStyle(TableStyle([
                            ('BACKGROUND',   (0,0), (-1,0), colors.HexColor('#1a365d')),
                            ('TEXTCOLOR',    (0,0), (-1,0), colors.white),
                            ('FONTNAME',     (0,0), (-1,0), 'Helvetica-Bold'),
                            ('FONTSIZE',     (0,0), (-1,0), 8.5),
                            ('ALIGN',        (0,0), (-1,-1), 'CENTER'),
                            ('VALIGN',       (0,0), (-1,-1), 'MIDDLE'),
                            ('FONTNAME',     (0,1), (-1,-1), 'Times-Roman'),
                            ('FONTSIZE',     (0,1), (-1,-1), 8.5),
                            ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.HexColor('#f5f7fa'), colors.white]),
                            ('GRID',         (0,0), (-1,-1), 0.5, colors.HexColor('#cccccc')),
                            ('TOPPADDING',   (0,0), (-1,-1), 5),
                            ('BOTTOMPADDING',(0,0), (-1,-1), 5),
                            ('LEFTPADDING',  (0,0), (-1,-1), 6),
                            ('RIGHTPADDING', (0,0), (-1,-1), 6),
                            ('BOX',          (0,0), (-1,-1), 1, colors.HexColor('#1a365d')),
                        ]))
                        story.append(tbl)
                        story.append(Spacer(1, 10))
                    continue
                if re.match(r'^(\d+\.|•|-|\*)\s', p):
                    story.append(Paragraph(p, styles['TextoLista']))
                elif clave == 'referencias':
                    story.append(Paragraph(p, styles['Referencia']))
                else:
                    story.append(Paragraph(p, styles['TextoJustificado']))
        else:
            story.append(Paragraph("Esta sección no pudo generarse.", styles['TextoJustificado']))

        story.append(PageBreak())

    doc.build(story)
    return filename, filepath

# ============================================================
# GENERAR WORD
# ============================================================
def generar_word(datos_usuario, secciones):
    nombre      = datos_usuario.get('nombre', 'Estudiante')
    autores_extra = datos_usuario.get('autores_extra', [])
    tema        = datos_usuario.get('tema', 'Tema')
    asignatura  = datos_usuario.get('asignatura', '')
    profesor    = datos_usuario.get('profesor', '')
    institucion = datos_usuario.get('institucion', '')
    ciudad      = datos_usuario.get('ciudad', '')
    fecha       = datos_usuario.get('fecha', datetime.now().strftime('%d/%m/%Y'))
    norma       = datos_usuario.get('norma', 'APA 7')

    doc = Document()
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

    # PORTADA
    portada = doc.add_paragraph()
    portada.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = portada.add_run('\n\n\nINFORME ACADÉMICO')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    doc.add_paragraph()
    p_tema = doc.add_paragraph()
    p_tema.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_tema.add_run(tema.upper())
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x2d, 0x4a, 0x7a)

    doc.add_paragraph()
    doc.add_paragraph()

    def agregar_dato(label, valor):
        if valor:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p.add_run(f"{label}: ")
            r1.bold = True
            r1.font.size = Pt(11)
            p.add_run(valor).font.size = Pt(11)

    agregar_dato("Presentado por", nombre)
    for autor in autores_extra:
        if autor.get('nombre'):
            linea = autor['nombre']
            if autor.get('cargo'):
                linea += f" — {autor['cargo']}"
            agregar_dato("Autor", linea)
    agregar_dato("Asignatura", asignatura)
    agregar_dato("Docente", profesor)
    agregar_dato("Institución", institucion)
    agregar_dato("Ciudad", ciudad)
    agregar_dato("Fecha", fecha)
    agregar_dato("Norma bibliográfica", norma)

    doc.add_page_break()

    # ÍNDICE
    h_idx = doc.add_heading('ÍNDICE', level=1)
    h_idx.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
    for item in ['1. Introducción', '2. Objetivos', '3. Marco Teórico', '4. Metodología',
                 '5. Desarrollo', '6. Conclusiones', '7. Recomendaciones', '8. Referencias Bibliográficas']:
        p = doc.add_paragraph()
        r = p.add_run(f"    {item}")
        r.font.size = Pt(11)

    doc.add_page_break()

    # SECCIONES
    secciones_orden = [
        ("1. INTRODUCCIÓN",               'introduccion'),
        ("2. OBJETIVOS",                  'objetivos'),
        ("3. MARCO TEÓRICO",              'marco_teorico'),
        ("4. METODOLOGÍA",                'metodologia'),
        ("5. DESARROLLO",                 'desarrollo'),
        ("6. CONCLUSIONES",               'conclusiones'),
        ("7. RECOMENDACIONES",            'recomendaciones'),
        ("8. REFERENCIAS BIBLIOGRÁFICAS", 'referencias'),
    ]

    for titulo_sec, clave in secciones_orden:
        h = doc.add_heading(titulo_sec, level=1)
        if h.runs:
            h.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
            h.runs[0].font.size = Pt(14)

        contenido_raw = secciones.get(clave, '')
        contenido = limpiar_para_word(contenido_raw)

        if contenido and len(contenido) > 50:
            # ── Renderizado especial para OBJETIVOS en Word ─────────
            if clave == 'objetivos':
                lineas = [l.strip() for l in re.split(r'\n+', contenido) if l.strip()]
                for linea in lineas:
                    if re.match(r'^OBJETIVO\s+(GENERAL|ESPEC)', linea, re.IGNORECASE):
                        p_obj = doc.add_paragraph()
                        r_obj = p_obj.add_run(linea)
                        r_obj.bold = True
                        r_obj.font.size = Pt(11)
                        r_obj.font.name = 'Times New Roman'
                        r_obj.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
                        p_obj.paragraph_format.space_before = Pt(10)
                        p_obj.paragraph_format.space_after  = Pt(3)
                    elif re.match(r'^\d+\.', linea):
                        p_obj = doc.add_paragraph()
                        r_obj = p_obj.add_run(linea)
                        r_obj.font.size = Pt(11)
                        r_obj.font.name = 'Times New Roman'
                        p_obj.paragraph_format.left_indent  = Pt(16)
                        p_obj.paragraph_format.space_before = Pt(6)
                        p_obj.paragraph_format.space_after  = Pt(6)
                        p_obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    else:
                        p_obj = doc.add_paragraph(linea)
                        p_obj.runs[0].font.size = Pt(11) if p_obj.runs else None
                        p_obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                doc.add_page_break()
                continue

            contenido_proc, tablas = extraer_tablas(contenido)
            parrafos = re.split(r'\n{2,}', contenido_proc)
            if len(parrafos) == 1:
                parrafos = contenido_proc.split('\n')
            for parrafo in parrafos:
                parrafo = parrafo.strip()
                if not parrafo:
                    continue
                # Detectar marcador de tabla
                m_tabla = re.match(r'^__TABLA_(\d+)__$', parrafo)
                if m_tabla and tablas:
                    idx_t = int(m_tabla.group(1))
                    if idx_t < len(tablas):
                        t = tablas[idx_t]
                        if t.get("titulo"):
                            p_titulo = doc.add_paragraph()
                            r_titulo = p_titulo.add_run(t["titulo"])
                            r_titulo.bold = True
                            r_titulo.italic = True
                            r_titulo.font.size = Pt(10)
                            r_titulo.font.name = 'Times New Roman'
                            r_titulo.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
                        cabeceras = t.get("cabeceras", [])
                        filas     = t.get("filas", [])
                        n_cols = max(len(cabeceras), max((len(f) for f in filas), default=0))
                        if n_cols == 0:
                            continue
                        def normalizar(fila, n):
                            return fila[:n] + [''] * max(0, n - len(fila))
                        tbl = doc.add_table(rows=1 + len(filas), cols=n_cols)
                        tbl.style = 'Table Grid'
                        # Cabecera
                        hdr_cells = tbl.rows[0].cells
                        for j, cab in enumerate(normalizar(cabeceras, n_cols)):
                            cell = hdr_cells[j]
                            cell.text = cab
                            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cab)
                            run.bold = True
                            run.font.size = Pt(9)
                            run.font.name = 'Calibri'
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            # Color de fondo cabecera
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:val'), 'clear')
                            shd.set(qn('w:color'), 'auto')
                            shd.set(qn('w:fill'), '1a365d')
                            tcPr.append(shd)
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Filas de datos
                        for i_f, fila in enumerate(filas):
                            row_cells = tbl.rows[i_f + 1].cells
                            bg_color = 'EEF2F7' if i_f % 2 == 0 else 'FFFFFF'
                            for j, val in enumerate(normalizar(fila, n_cols)):
                                cell = row_cells[j]
                                cell.text = val
                                run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(val)
                                run.font.size = Pt(9)
                                run.font.name = 'Calibri'
                                # Color alterno
                                tc = cell._tc
                                tcPr = tc.get_or_add_tcPr()
                                shd = OxmlElement('w:shd')
                                shd.set(qn('w:val'), 'clear')
                                shd.set(qn('w:color'), 'auto')
                                shd.set(qn('w:fill'), bg_color)
                                tcPr.append(shd)
                                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph()
                    continue
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                r = p.add_run(parrafo)
                if titulo_sec == '8. REFERENCIAS BIBLIOGRÁFICAS':
                    # Sangría francesa APA: primera línea al margen, resto indentado
                    r.font.size = Pt(10.5)
                    r.font.name = 'Times New Roman'
                    p.paragraph_format.left_indent       = Cm(1.0)
                    p.paragraph_format.first_line_indent = Cm(-1.0)
                    p.paragraph_format.space_after        = Pt(6)
                else:
                    r.font.size = Pt(11)
                    r.font.name = 'Times New Roman'
                    p.paragraph_format.first_line_indent = Cm(0.5)
                    p.paragraph_format.space_after = Pt(8)
        else:
            p = doc.add_paragraph("Esta sección no pudo generarse correctamente.")
            p.runs[0].font.size = Pt(11)

        doc.add_page_break()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ============================================================
# RUTAS — PÁGINAS (GET)
# ============================================================
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generar', methods=['GET'])
def generar_page():
    return render_template('generar.html')

@app.route('/auth')
def auth_page():
    return render_template('auth.html')

@app.route('/mis-informes')
def mis_informes():
    return render_template('mis-informes.html')

@app.route('/perfil')
def perfil():
    return render_template('perfil.html')

# ============================================================
# RUTAS — API (POST)
# ============================================================
@app.route('/api/generar', methods=['POST'])
@limit("15 per hour")   # máx 15 generaciones por IP por hora (protege la API de DeepSeek)
@requiere_auth
def api_generar(user_id):
    try:
        data = request.get_json(silent=True) or {}

        # Validar parámetros de negocio
        params_ok, params_err = validar_params_informe(data)
        if not params_ok:
            return jsonify({'success': False, 'error': params_err}), 400

        tema            = sanitizar_texto(data.get('tema', '').strip(), 500)
        nivel           = data.get('nivel', 'universitario')
        tipo_informe    = data.get('tipo_informe', 'academico')
        norma           = data.get('norma', 'APA 7')
        nombre          = sanitizar_nombre(data.get('nombre', 'Estudiante'))
        asignatura      = sanitizar_texto(data.get('asignatura', ''), 200)
        profesor        = sanitizar_nombre(data.get('profesor', ''))
        institucion     = sanitizar_texto(data.get('institucion', ''), 200)
        ciudad          = sanitizar_texto(data.get('ciudad', ''), 100)
        texto_usuario   = sanitizar_texto(data.get('texto_usuario', ''), 10000)
        autores         = data.get('autores', [])
        if not isinstance(autores, list):
            autores = []
        nombre_principal = sanitizar_nombre(autores[0].get('nombre', nombre)) if autores else nombre

        logger.info(f"📨 Generando informe — Tema: {tema[:50]}... | Usuario: {user_id}")
        secciones, refs_reales = generar_informe_completo(tema, texto_usuario, tipo_informe, norma, nivel, modo=data.get('modo', 'rapido'))

        if not secciones:
            return jsonify({'success': False, 'error': 'No se pudo generar el informe'}), 500

        datos_usuario = {
            'nombre':        nombre_principal,
            'autores_extra': autores[1:] if len(autores) > 1 else [],
            'tema':          tema,
            'asignatura':    asignatura,
            'profesor':      profesor,
            'institucion':   institucion,
            'ciudad':        ciudad,
            'fecha':         datetime.now().strftime('%d/%m/%Y'),
            'norma':         norma
        }

        if DB_DISPONIBLE:
            try:
                guardar_informe(
                    user_id=user_id,
                    datos_usuario=datos_usuario,
                    secciones=secciones,
                    refs_reales=refs_reales,
                    tipo_informe=tipo_informe,
                    norma=norma,
                    nivel=nivel,
                    modo=data.get('modo', 'rapido')
                )
                logger.info(f"Informe guardado para usuario {user_id}")
            except Exception as e:
                logger.error(f"Error guardando informe: {e}")

        return jsonify({'success': True, 'secciones': secciones, 'datos_usuario': datos_usuario})

    except Exception as e:
        logger.error(f"Error en /api/generar: {e}")
        return jsonify({'success': False, 'error': 'Error interno al generar el informe'}), 500


@app.route('/api/generar-seccion', methods=['POST'])
@limit("30 per hour")
@requiere_auth
def api_generar_seccion(user_id):
    try:
        data      = request.get_json(silent=True) or {}
        seccion   = sanitizar_texto(data.get('seccion', ''), 50)
        tema      = sanitizar_texto(data.get('tema', '').strip(), 500)
        nivel     = data.get('nivel', 'universitario')
        tipo      = data.get('tipo_informe', 'academico')
        norma     = data.get('norma', 'APA 7')
        info      = sanitizar_texto(data.get('texto_usuario', ''), 10000)
        refs_man  = sanitizar_texto(data.get('refs_manuales', ''), 5000)
        objetivos_texto = sanitizar_texto(data.get('objetivos_texto', ''), 2000)

        # Validar enumeraciones
        if nivel not in NIVELES_VALIDOS:
            nivel = 'universitario'
        if tipo not in TIPOS_VALIDOS:
            tipo = 'academico'
        if norma not in NORMAS_VALIDAS:
            norma = 'APA 7'

        if not seccion or not tema:
            return jsonify({'success': False, 'error': 'Faltan parámetros'}), 400

        modo      = data.get('modo', 'rapido')
        if modo not in MODOS_VALIDOS:
            modo = 'rapido'

        contenido = generar_seccion(seccion, tema, info, tipo, norma, nivel, refs_man, modo, objetivos_texto,
                                    contexto_refs=data.get('contexto_refs', ''))

        if contenido:
            return jsonify({'success': True, 'seccion': seccion, 'contenido': contenido})
        return jsonify({'success': False, 'error': f'No se pudo generar: {seccion}'}), 500

    except Exception as e:
        logger.error(f"Error en /api/generar-seccion: {e}")
        return jsonify({'success': False, 'error': 'Error interno'}), 500


@app.route('/api/referencias-reales', methods=['POST'])
def api_referencias_reales():
    """
    Obtiene referencias REALES de CrossRef y OpenAlex,
    las formatea según la norma y las devuelve al frontend.
    """
    try:
        data         = request.json
        tema         = data.get('tema', '').strip()
        norma        = data.get('norma', 'APA 7')
        refs_manuales = data.get('refs_manuales', '')

        if not tema:
            return jsonify({'success': False, 'error': 'Tema requerido'}), 400

        logger.info(f"🔬 Buscando referencias reales para: '{tema[:50]}'")

        # Extraer contenido relevante del informe para mejorar la relevancia
        # Usamos introducción + marco teórico + desarrollo si están disponibles
        contenido_informe = data.get('contenido_informe', '')
        if not contenido_informe:
            # Intentar reconstruir desde secciones si el frontend las envía
            secciones_recibidas = data.get('secciones', {})
            partes = []
            for sec in ('introduccion', 'marco_teorico', 'desarrollo'):
                if secciones_recibidas.get(sec):
                    partes.append(secciones_recibidas[sec][:1500])
            contenido_informe = ' '.join(partes)

        # Buscar en CrossRef + OpenAlex, usando el contenido para mejorar relevancia
        refs = buscar_referencias_reales(tema, cantidad_total=12,
                                         contenido_informe=contenido_informe)

        if not refs:
            # Fallback: usar DeepSeek para generar referencias plausibles
            logger.warning("No se encontraron referencias reales, usando fallback IA")
            contenido_ia = generar_seccion('referencias', tema, refs_manuales, 'academico', norma, 'universitario', refs_manuales)
            return jsonify({
                'success':    True,
                'contenido':  contenido_ia or 'No se pudieron obtener referencias.',
                'fuente':     'ia_fallback',
                'total':      0
            })

        # Formatear según la norma
        texto_refs = formatear_referencias(refs, norma)

        # Si hay referencias manuales, agregarlas al final
        if refs_manuales and refs_manuales.strip():
            texto_refs += f"\n\n{refs_manuales.strip()}"

        logger.info(f"✅ {len(refs)} referencias reales formateadas en norma {norma}")

        return jsonify({
            'success':   True,
            'contenido': texto_refs,
            'fuente':    'crossref_openalex',
            'total':     len(refs),
            'detalles':  [{'tipo': r['tipo'], 'titulo': r['titulo'][:60], 'anio': r['anio']} for r in refs]
        })

    except Exception as e:
        logger.error(f"Error en /api/referencias-reales: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/exportar-pdf', methods=['POST'])
def exportar_pdf():
    try:
        data = request.json
        filename, filepath = generar_pdf(data['datos_usuario'], data['secciones'])
        return send_file(filepath, as_attachment=True, download_name=filename, mimetype='application/pdf')
    except Exception as e:
        logger.error(f"Error PDF: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/exportar-word', methods=['POST'])
def exportar_word():
    try:
        data   = request.json
        buffer = generar_word(data['datos_usuario'], data['secciones'])
        nombre_archivo = f"informe_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        return send_file(
            buffer,
            as_attachment=True,
            download_name=nombre_archivo,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        logger.error(f"Error Word: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================================
# RUTAS DE AUTENTICACIÓN Y PERFIL (CAMBIO 3)
# ============================================================
@app.route('/api/auth/registro', methods=['POST'])
@limit("5 per hour")   # máx 5 registros por IP por hora (anti-spam)
def api_registro():
    """Registrar un nuevo usuario con validaciones de seguridad."""
    try:
        data = request.get_json(silent=True) or {}
        nombre   = sanitizar_nombre(data.get('nombre', ''))
        email    = sanitizar_texto(data.get('email', ''), 254).lower()
        password = data.get('password', '')

        # Validar email
        if not es_email_valido(email):
            return jsonify({'success': False, 'error': 'El formato del correo no es válido'}), 400

        # Validar contraseña con requisitos de seguridad
        pwd_ok, pwd_msg = es_password_seguro(password)
        if not pwd_ok:
            return jsonify({'success': False, 'error': pwd_msg}), 400

        if not nombre:
            return jsonify({'success': False, 'error': 'El nombre es requerido'}), 400

        resultado = registrar_usuario(nombre, email, password)
        if resultado['success']:
            log_evento_seguridad('REGISTRO_EXITOSO', f"email={email}", request)
            return jsonify({'success': True, 'user_id': resultado['user_id'],
                            'auto_login': resultado.get('auto_login', False),
                            'access_token': resultado.get('access_token', '')})
        else:
            log_evento_seguridad('REGISTRO_FALLIDO', f"email={email}", request)
            return jsonify({'success': False, 'error': resultado['error']}), 400

    except Exception as e:
        logger.error(f"Error en registro: {e}")
        return jsonify({'success': False, 'error': 'Error interno. Intenta de nuevo.'}), 500


@app.route('/api/auth/login', methods=['POST'])
@limit("10 per hour")   # máx 10 intentos por IP por hora (anti-brute-force)
def api_login():
    """Iniciar sesión — devuelve el access_token JWT."""
    try:
        data = request.get_json(silent=True) or {}
        email    = sanitizar_texto(data.get('email', ''), 254).lower()
        password = data.get('password', '')

        if not email or not password:
            return jsonify({'success': False, 'error': 'Email y contraseña son requeridos'}), 400

        if not es_email_valido(email):
            # Mismo mensaje que credenciales incorrectas para no revelar info
            return jsonify({'success': False, 'error': 'Email o contraseña incorrectos'}), 401

        resultado = login_usuario(email, password)
        if resultado['success']:
            log_evento_seguridad('LOGIN_EXITOSO', f"email={email}", request)
            return jsonify({
                'success':      True,
                'user_id':      resultado['user_id'],
                'nombre':       resultado['nombre'],
                'email':        resultado['email'],
                'access_token': resultado.get('access_token', ''),  # JWT para verificar en el backend
            })
        else:
            log_evento_seguridad('LOGIN_FALLIDO', f"email={email}", request)
            # Mensaje genérico: no decir si el email existe o no
            return jsonify({'success': False, 'error': 'Email o contraseña incorrectos'}), 401

    except Exception as e:
        logger.error(f"Error en login: {e}")
        return jsonify({'success': False, 'error': 'Error interno. Intenta de nuevo.'}), 500


@app.route('/api/perfil', methods=['GET'])
@requiere_auth
def api_obtener_perfil(user_id):
    """Obtener perfil de usuario — requiere JWT válido."""
    try:
        perfil = obtener_perfil(user_id)
        if perfil:
            return jsonify({'success': True, 'perfil': perfil})
        else:
            return jsonify({'success': False, 'error': 'Usuario no encontrado'}), 404
    except Exception as e:
        logger.error(f"Error obteniendo perfil: {e}")
        return jsonify({'success': False, 'error': 'Error interno'}), 500


@app.route('/api/perfil', methods=['PUT'])
@requiere_auth
def api_actualizar_perfil(user_id):
    """Actualizar perfil de usuario — requiere JWT válido."""
    try:
        data = request.get_json(silent=True) or {}
        # Sanitizar cada campo antes de guardarlo
        datos_limpios = {
            'nombre':      sanitizar_nombre(data.get('nombre', '')),
            'institucion': sanitizar_texto(data.get('institucion', ''), 200),
            'carrera':     sanitizar_texto(data.get('carrera', ''), 200),
            'ciudad':      sanitizar_texto(data.get('ciudad', ''), 100),
            'telefono':    re.sub(r'[^\d\+\-\s\(\)]', '', data.get('telefono', ''))[:20],
        }
        resultado = actualizar_perfil(user_id, datos_limpios)
        if resultado['success']:
            return jsonify({'success': True})
        else:
            return jsonify({'success': False, 'error': resultado['error']}), 400
    except Exception as e:
        logger.error(f"Error actualizando perfil: {e}")
        return jsonify({'success': False, 'error': 'Error interno'}), 500


@app.route('/api/mis-informes', methods=['GET'])
@requiere_auth
def api_mis_informes(user_id):
    """Obtener lista de informes del usuario — requiere JWT válido."""
    try:
        limit_q  = min(request.args.get('limit', 50, type=int), 100)   # máx 100
        offset_q = max(request.args.get('offset', 0, type=int), 0)
        informes = obtener_mis_informes(user_id, limit_q, offset_q)
        return jsonify({'success': True, 'informes': informes, 'total': len(informes)})
    except Exception as e:
        logger.error(f"Error obteniendo informes: {e}")
        return jsonify({'success': False, 'error': 'Error interno'}), 500


@app.route('/api/informe/<informe_id>', methods=['GET'])
@requiere_auth
def api_obtener_informe(user_id, informe_id):
    """Obtener un informe específico — requiere JWT válido y que sea del usuario."""
    try:
        if not es_uuid_valido(informe_id):
            return jsonify({'success': False, 'error': 'ID inválido'}), 400
        informe = obtener_informe(informe_id, user_id)
        if informe:
            return jsonify({'success': True, 'informe': informe})
        else:
            # No revelar si existe pero es de otro usuario vs. no existe
            return jsonify({'success': False, 'error': 'Informe no encontrado'}), 404
    except Exception as e:
        logger.error(f"Error obteniendo informe: {e}")
        return jsonify({'success': False, 'error': 'Error interno'}), 500


@app.route('/api/informe/<informe_id>', methods=['DELETE'])
@requiere_auth
def api_eliminar_informe(user_id, informe_id):
    """Eliminar un informe — requiere JWT válido y que sea del usuario."""
    try:
        if not es_uuid_valido(informe_id):
            return jsonify({'success': False, 'error': 'ID inválido'}), 400
        resultado = eliminar_informe(informe_id, user_id)
        if resultado['success']:
            log_evento_seguridad('INFORME_ELIMINADO', f"user={user_id} informe={informe_id}", request)
            return jsonify({'success': True})
        else:
            return jsonify({'success': False, 'error': resultado['error']}), 400
    except Exception as e:
        logger.error(f"Error eliminando informe: {e}")
        return jsonify({'success': False, 'error': 'Error interno'}), 500

@app.route('/api/estadisticas', methods=['GET'])
@requiere_auth
def api_estadisticas(user_id):
    """
    Devuelve estadísticas agregadas del usuario autenticado.
 
    Respuesta ejemplo:
    {
        "success": true,
        "stats": {
            "total_informes":     8,
            "total_referencias":  62,
            "norma_mas_usada":    "APA 7",
            "tipo_mas_usado":     "academico",
            "nivel_mas_usado":    "universitario",
            "dias_activo":        5,
            "primer_informe":     "2025-01-10T...",
            "ultimo_informe":     "2025-05-01T...",
            "informes_por_norma": {"APA 7": 5, "ICONTEC": 3},
            "informes_por_tipo":  {"academico": 6, "tesis": 2},
            "racha_actual":       2
        }
    }
    """
    stats = obtener_estadisticas_usuario(user_id)
    if stats is None:
        return jsonify({'success': False, 'error': 'No se pudieron obtener estadísticas'}), 500
    return jsonify({'success': True, 'stats': stats})
 
 
@app.route('/api/estadisticas/actividad', methods=['GET'])
@requiere_auth
def api_actividad(user_id):
    """
    Devuelve la actividad diaria del usuario en los últimos N días.
    Query param opcional: ?dias=30  (default 30, máximo 90)
 
    Respuesta ejemplo:
    {
        "success": true,
        "actividad": [
            {"fecha": "2025-04-28", "cantidad": 1},
            {"fecha": "2025-05-01", "cantidad": 3}
        ],
        "dias": 30
    }
    """
    try:
        dias = min(int(request.args.get('dias', 30)), 90)
    except (ValueError, TypeError):
        dias = 30
 
    actividad = obtener_resumen_actividad(user_id, dias=dias)
    return jsonify({'success': True, 'actividad': actividad, 'dias': dias})

@app.route('/health')
def health():
    # En producción no revelar detalles internos
    env = os.environ.get('FLASK_ENV', 'production')
    if env == 'production':
        return jsonify({'status': 'healthy', 'version': '3.3'})
    return jsonify({
        'status':              'healthy',
        'api_configured':      bool(DEEPSEEK_API_KEY),
        'normas_disponibles':  list(NORMAS_INSTRUCCIONES.keys()),
        'tipos_disponibles':   list(TIPOS_INSTRUCCIONES.keys()),
        'version':             '3.3',
        'database_available':  DB_DISPONIBLE
    })




# ─────────────────────────────────────────────────────────────
# FEEDBACK / RECOMENDACIONES  →  envío por email al dueño
# ─────────────────────────────────────────────────────────────
@app.route('/api/feedback', methods=['POST'])
@limit("5 per hour")   # máx 5 feedbacks por IP por hora
def api_feedback():
    """Recibe una sugerencia y la envía al correo del dueño."""
    try:
        data    = request.get_json(silent=True) or {}
        nombre  = sanitizar_nombre(data.get('nombre', 'Anónimo'))
        correo  = sanitizar_texto(data.get('correo', ''), 254)
        tipo    = sanitizar_texto(data.get('tipo', 'sugerencia'), 50)
        mensaje = sanitizar_texto(data.get('mensaje', ''), 1000)

        # Validar correo si se proporcionó
        if correo and not es_email_valido(correo):
            correo = ''  # descartar silenciosamente (no es campo requerido)

        # Validar tipo contra lista permitida
        tipos_permitidos = {'sugerencia', 'error', 'funcion', 'otro'}
        if tipo not in tipos_permitidos:
            tipo = 'sugerencia'

        if not mensaje or len(mensaje) < 10:
            return jsonify({'success': False, 'error': 'El mensaje es demasiado corto'}), 400

        import smtplib
        from email.mime.text import MIMEText
        from email.mime.multipart import MIMEMultipart

        SMTP_HOST  = os.environ.get('SMTP_HOST', 'smtp.gmail.com')
        SMTP_PORT  = int(os.environ.get('SMTP_PORT', '587'))
        SMTP_USER  = os.environ.get('SMTP_USER', '')
        SMTP_PASS  = os.environ.get('SMTP_PASS', '')
        OWNER_MAIL = os.environ.get('OWNER_EMAIL', SMTP_USER)

        if not SMTP_USER or not SMTP_PASS:
            logger.warning(f"[FEEDBACK sin SMTP] {nombre} | {correo} | {tipo}: {mensaje[:80]}")
            return jsonify({'success': True})

        # ✅ SEGURO: usar sanitizar_html_email en todos los campos del usuario
        nombre_safe  = sanitizar_html_email(nombre)
        correo_safe  = sanitizar_html_email(correo) if correo else 'No proporcionado'
        tipo_safe    = sanitizar_html_email(tipo)
        mensaje_safe = sanitizar_html_email(mensaje)

        asunto = f"[ARP Feedback] {tipo.capitalize()} de {nombre}"
        cuerpo = f"""
<h2>Nueva sugerencia — Academic Report Pro</h2>
<table style="border-collapse:collapse;font-family:sans-serif;font-size:14px">
  <tr><td style="padding:6px 14px;font-weight:bold;color:#888">Tipo</td>
      <td style="padding:6px 14px">{tipo_safe}</td></tr>
  <tr><td style="padding:6px 14px;font-weight:bold;color:#888">Nombre</td>
      <td style="padding:6px 14px">{nombre_safe}</td></tr>
  <tr><td style="padding:6px 14px;font-weight:bold;color:#888">Correo</td>
      <td style="padding:6px 14px">{correo_safe}</td></tr>
</table>
<h3 style="margin-top:20px">Mensaje</h3>
<div style="background:#f5f5f5;padding:16px;border-radius:8px;font-size:14px;line-height:1.6">
  {mensaje_safe}
</div>
<p style="font-size:12px;color:#aaa;margin-top:20px">
  Enviado desde Academic Report Pro · {datetime.now().strftime('%d/%m/%Y %H:%M')}
</p>
"""
        msg = MIMEMultipart('alternative')
        msg['Subject'] = asunto
        msg['From']    = SMTP_USER
        msg['To']      = OWNER_MAIL
        # Solo agregar Reply-To si el correo es válido
        if correo and es_email_valido(correo):
            msg['Reply-To'] = correo
        msg.attach(MIMEText(cuerpo, 'html', 'utf-8'))

        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
            server.starttls()
            server.login(SMTP_USER, SMTP_PASS)
            server.sendmail(SMTP_USER, OWNER_MAIL, msg.as_string())

        logger.info(f"Feedback enviado de {nombre} ({correo}): {tipo}")
        return jsonify({'success': True})

    except Exception as e:
        logger.error(f"Error enviando feedback: {e}")
        return jsonify({'success': True})   # no revelar errores SMTP al usuario


# ─────────────────────────────────────────────────────────────
# RECUPERACIÓN DE CONTRASEÑA  (Supabase lo maneja)
# ─────────────────────────────────────────────────────────────
@app.route('/api/auth/verificar-token', methods=['POST'])
@limit("10 per hour")
def api_verificar_token():
    """
    Verifica un access_token de Supabase y devuelve los datos del usuario.
    Se usa cuando el usuario llega desde el link de verificación de correo.
    """
    try:
        from database import supabase, DB_DISPONIBLE
        data  = request.get_json(silent=True) or {}
        token = data.get('access_token', '').strip()

        if not token:
            return jsonify({'success': False, 'error': 'Token requerido'}), 400

        if not DB_DISPONIBLE or supabase is None:
            return jsonify({'success': False, 'error': 'Base de datos no disponible'}), 503

        # Verificar el token con Supabase
        user_response = supabase.auth.get_user(token)
        if not user_response or not user_response.user:
            return jsonify({'success': False, 'error': 'Token inválido o expirado'}), 401

        user   = user_response.user
        nombre = ''
        if user.user_metadata:
            nombre = user.user_metadata.get('nombre', '')

        log_evento_seguridad('VERIFICACION_CORREO', f"email={user.email}", request)
        return jsonify({
            'success':  True,
            'user_id':  str(user.id),
            'email':    user.email,
            'nombre':   nombre,
        })

    except Exception as e:
        logger.error(f"Error verificando token: {e}")
        return jsonify({'success': False, 'error': 'Token inválido'}), 401
def api_recuperar_password():
    """Envía email de recuperación a través de Supabase."""
    try:
        from database import supabase, DB_DISPONIBLE
        data  = request.get_json(silent=True) or {}
        email = sanitizar_texto(data.get('email', ''), 254).lower()

        if not email or not es_email_valido(email):
            # Siempre éxito: no revelar si el email existe
            return jsonify({'success': True})

        if not DB_DISPONIBLE or supabase is None:
            return jsonify({'success': True})   # silencioso para no revelar estado del sistema

        redirect_url = data.get('redirect_url', '')
        # Validar que el redirect_url sea del propio dominio
        dominio_permitido = os.environ.get('SITE_URL', '')
        if redirect_url and dominio_permitido and not redirect_url.startswith(dominio_permitido):
            redirect_url = dominio_permitido

        opts = {'redirect_to': redirect_url} if redirect_url else {}
        supabase.auth.reset_password_email(email, opts)
        log_evento_seguridad('RECUPERAR_PASSWORD', f"email={email}", request)
        return jsonify({'success': True})
    except Exception as e:
        logger.error(f"Error recuperando contraseña: {e}")
        return jsonify({'success': True})   # siempre éxito para no revelar si el email existe



@app.route('/api/refs-previas', methods=['POST'])
def api_refs_previas():
    try:
        data  = request.get_json(silent=True) or {}
        tema  = data.get('tema', '').strip()
        norma = data.get('norma', 'APA 7')
        if not tema:
            return jsonify({'success': False, 'error': 'Tema requerido'}), 400
        logger.info(f"Buscando refs previas para: '{tema[:50]}'")
        refs = buscar_referencias_reales(tema, cantidad_total=12)
        if not refs:
            return jsonify({'success': True, 'contexto_refs': '', 'total': 0, 'refs': []})
        contexto_refs    = _construir_contexto_refs(refs, norma)
        refs_formateadas = formatear_referencias(refs, norma)
        logger.info(f"Refs previas obtenidas: {len(refs)}")
        return jsonify({
            'success':       True,
            'contexto_refs': contexto_refs,
            'refs_texto':    refs_formateadas,
            'total':         len(refs),
            'refs':          [{'tipo': r['tipo'], 'titulo': r['titulo'][:70],
                               'anio': r['anio'], 'doi': r.get('doi', '')} for r in refs]
        })
    except Exception as e:
        logger.error(f"Error en /api/refs-previas: {e}")
        return jsonify({'success': True, 'contexto_refs': '', 'total': 0, 'refs': []})


@app.route('/api/guardar-informe', methods=['POST'])
@requiere_auth
def api_guardar_informe(user_id):
    try:
        if not DB_DISPONIBLE:
            return jsonify({'success': False, 'error': 'Base de datos no disponible'}), 503
        data          = request.get_json(silent=True) or {}
        secciones     = data.get('secciones', {})
        datos_usuario = data.get('datos_usuario', {})
        refs_reales   = data.get('refs_reales', [])
        tipo_informe  = data.get('tipo_informe', 'academico')
        norma         = data.get('norma', 'APA 7')
        nivel         = data.get('nivel', 'universitario')
        modo          = data.get('modo', 'rapido')

        # Validar enumeraciones
        if tipo_informe not in TIPOS_VALIDOS: tipo_informe = 'academico'
        if norma not in NORMAS_VALIDAS: norma = 'APA 7'
        if nivel not in NIVELES_VALIDOS: nivel = 'universitario'
        if modo not in MODOS_VALIDOS: modo = 'rapido'

        if not secciones or not datos_usuario.get('tema'):
            return jsonify({'success': False, 'error': 'Datos incompletos'}), 400
        informe_id = guardar_informe(
            user_id=user_id,
            datos_usuario=datos_usuario,
            secciones=secciones,
            refs_reales=refs_reales,
            tipo_informe=tipo_informe,
            norma=norma,
            nivel=nivel,
            modo=modo
        )
        if informe_id:
            logger.info(f"Informe guardado: {informe_id}")
            return jsonify({'success': True, 'informe_id': informe_id})
        return jsonify({'success': False, 'error': 'No se pudo guardar'}), 500
    except Exception as e:
        logger.error(f"Error en /api/guardar-informe: {e}")
        return jsonify({'success': False, 'error': 'Error interno'}), 500


@app.route('/recuperar')
def recuperar_page():
    return render_template('recuperar.html')

@app.route('/sugerencias')
def sugerencias_page():
    return render_template('sugerencias.html')

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 10000))
    app.run(host='0.0.0.0', port=port, debug=False)
